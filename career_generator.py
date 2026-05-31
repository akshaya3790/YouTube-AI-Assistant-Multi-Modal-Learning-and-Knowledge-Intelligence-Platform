import json
import io
from datetime import datetime
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage
from fpdf import FPDF
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_career_report(transcript, title, chapters, category, selected_model="gemini-2.5-flash"):
    """
    Queries Gemini to parse video transcript and chapters for career intelligence data.
    """
    chapters_text = ""
    if chapters:
        for i, ch in enumerate(chapters):
            chapters_text += f"- Chapter {i+1}: {ch.get('title', 'Chapter')} ({ch.get('start_time', '00:00')}) - Summary: {ch.get('summary', '')}\n"

    prompt = f"""
    You are a Senior Career Intelligence Architect, Resume Optimization Expert, and ATS Specialist.
    Analyze the transcript and chapter info for the video "{title}" in the domain category "{category}" and extract career intelligence details.
    
    Video Chapters Breakdown:
    {chapters_text}
    
    Transcript snippet:
    {transcript[:28000]}
    
    Instructions:
    1. Parse the technical/professional skills learned from this video. For each skill, assign an Importance Score (1-100), Industry Demand Score (1-100), and Resume Relevance Score (1-100).
    2. Extract all technologies mentioned. Detail their Category (e.g. Programming Language, Library, Platform), Usage Context, and Learning Level.
    3. Extract frameworks, tools, and industry terminology concepts.
    4. Generate 5-10 ATS-friendly keywords with Score Contribution weights and Industry Demand.
    5. Draft 3-5 high-impact, professional resume bullet points showcasing video learning achievements using strong action verbs (e.g., "Leveraged...", "Engineered...").
    6. Generate LinkedIn profile optimizations: Suggested Skills, Headline variations, and "About" section summaries.
    7. Generate 3-5 technical/scenario interview questions based on the video contents, with comprehensive answers.
    8. Suggest 2-3 suitable career roles with Match Scores based on the video context.
    9. Complete a Chapter-Wise Skill Extraction: for each chapter listed, identify the skills, technologies, frameworks, and tools discussed in that specific chapter.
    
    Return ONLY a raw, valid JSON object following the schema below. 
    Do NOT include markdown syntax (like ```json or ```), backticks, comments, or any other surrounding text.

    JSON Schema:
    {{
      "skills_learned": [
        {{
          "name": "Skill Name",
          "description": "Brief description of the skill.",
          "importance_score": 85,
          "demand_score": 90,
          "resume_relevance_score": 95
        }}
      ],
      "technologies": [
        {{
          "name": "Python",
          "category": "Programming Language",
          "learning_level": "Intermediate",
          "usage_context": "Used for data cleaning and scripting.",
          "mentioned_count": 12
        }}
      ],
      "frameworks": [
        {{
          "name": "TensorFlow",
          "purpose": "Deep learning models.",
          "industry_usage": "High in enterprise machine learning.",
          "resume_worthiness": 90
        }}
      ],
      "tools": [
        {{
          "name": "Docker",
          "category": "Containerization",
          "usage_context": "Packaging and deployment.",
          "skill_level": "Intermediate"
        }}
      ],
      "concepts": [
        {{
          "name": "Supervised Learning",
          "explanation": "Modeling with labeled data."
        }}
      ],
      "ats_keywords": [
        {{
          "keyword": "Predictive Modeling",
          "score_contribution": 15,
          "demand_score": 90
        }}
      ],
      "resume_bullets": [
        "Analyzed machine learning workflows in Python and Scikit-Learn to build and evaluate predictive models."
      ],
      "linkedin_suggestions": {{
        "skills": ["Machine Learning", "Python"],
        "headline_suggestions": ["AI Practitioner | Python Developer"],
        "about_additions": "Familiar with end-to-end data pipelines...",
        "suggested_certifications": ["Google TensorFlow Certification"]
      }},
      "interview_preparation": [
        {{
          "question": "What is overfitting?",
          "answer": "When a model learns noise in training data instead of general patterns.",
          "difficulty": "Intermediate"
        }}
      ],
      "recommended_careers": [
        {{
          "title": "Machine Learning Engineer",
          "match_score": 90
        }}
      ],
      "chapter_wise_skills": [
        {{
          "chapter_title": "Data Preprocessing",
          "timestamp": "02:40",
          "skills": ["Data Preprocessing"],
          "technologies": ["Python"],
          "frameworks": ["Pandas"],
          "tools": ["Jupyter"]
        }}
      ]
    }}
    """
    
    try:
        model = ChatGoogleGenerativeAI(model=selected_model)
        response = model.invoke([HumanMessage(content=prompt)])
        clean_text = response.content.replace("```json", "").replace("```", "").strip()
        
        start_idx = clean_text.find("{")
        end_idx = clean_text.rfind("}")
        if start_idx != -1 and end_idx != -1:
            clean_text = clean_text[start_idx:end_idx+1]
            
        report = json.loads(clean_text)
        return report
    except Exception as e:
        print(f"Error extracting career report: {e}")
        return get_fallback_career_report(title)

def get_fallback_career_report(title):
    return {
        "skills_learned": [
            {
                "name": "Core Technical Principles",
                "description": f"Understanding technical foundations demonstrated in {title}.",
                "importance_score": 80,
                "demand_score": 85,
                "resume_relevance_score": 90
            }
        ],
        "technologies": [
            {
                "name": "General Technology Stack",
                "category": "Development Framework",
                "learning_level": "Beginner",
                "usage_context": "Core programming workflows.",
                "mentioned_count": 5
            }
        ],
        "frameworks": [],
        "tools": [],
        "concepts": [
            {
                "name": "Systems Architecture",
                "explanation": "Understanding how software systems coordinate."
            }
        ],
        "ats_keywords": [
            {
                "keyword": "Technical Development",
                "score_contribution": 10,
                "demand_score": 80
            }
        ],
        "resume_bullets": [
            f"Studied software architectures and programming principles outlined in the tutorial on {title}."
        ],
        "linkedin_suggestions": {
            "skills": ["Software Development"],
            "headline_suggestions": ["Technology Enthusiast"],
            "about_additions": "Familiar with modern architecture paradigms.",
            "suggested_certifications": []
        },
        "interview_preparation": [
            {
                "question": "What is the primary topic of the video?",
                "answer": "The video outlines technical workflows, system integration, and implementation details.",
                "difficulty": "Beginner"
            }
        ],
        "recommended_careers": [
            {
                "title": "Software Developer",
                "match_score": 80
            }
        ],
        "chapter_wise_skills": []
    }

def generate_skill_gap_analysis(learned_skills, target_role, selected_model="gemini-2.5-flash"):
    """
    Queries Gemini to match user's learned skills against a target role and output a roadmap.
    """
    prompt = f"""
    You are a Career Path Consultant and Technical Recruiter.
    Perform a professional Skill Gap Analysis.
    
    Target Job Role: {target_role}
    Learned Skills: {', '.join(learned_skills)}
    
    Instructions:
    1. Evaluate the Match Score (0 to 100) based on how well the learned skills cover the responsibilities of the target job role.
    2. Identify the critical missing skills or technologies required for this role.
    3. Generate a structured step-by-step learning roadmap to bridge the gap.
    
    Return ONLY a raw, valid JSON object following the schema below.
    Do NOT include markdown syntax (like ```json or ```), backticks, comments, or any other surrounding text.

    JSON Schema:
    {{
      "match_score": 75,
      "missing_skills": [
        {{
          "name": "SQL Databases",
          "why_needed": "Required to query and manage relational datasets in production."
        }}
      ],
      "learning_roadmap": [
        "1. Learn database theory and basic SELECT queries.",
        "2. Build SQL projects with multi-table joins and indexes."
      ]
    }}
    """
    try:
        model = ChatGoogleGenerativeAI(model=selected_model)
        response = model.invoke([HumanMessage(content=prompt)])
        clean_text = response.content.replace("```json", "").replace("```", "").strip()
        
        start_idx = clean_text.find("{")
        end_idx = clean_text.rfind("}")
        if start_idx != -1 and end_idx != -1:
            clean_text = clean_text[start_idx:end_idx+1]
            
        return json.loads(clean_text)
    except Exception as e:
        print(f"Error generating skill gap analysis: {e}")
        return {
            "match_score": 50,
            "missing_skills": [
                {
                    "name": "Domain Specific Stack",
                    "why_needed": "Essential for backend/frontend operations."
                }
            ],
            "learning_roadmap": [
                "1. Research industry standard tools for target role.",
                "2. Implement core projects."
            ]
        }

def create_career_markdown(report, title):
    md = f"# AI Career Intelligence & Resume Report: {title}\n\n"
    
    md += "## 📊 Extracted Technical Skills\n"
    for s in report.get("skills_learned", []):
        md += f"### 💡 {s.get('name')}\n"
        md += f"- **Description**: {s.get('description')}\n"
        md += f"- **Importance**: {s.get('importance_score')}/100 | **Industry Demand**: {s.get('demand_score')}/100 | **Resume Relevance**: {s.get('resume_relevance_score')}/100\n\n"
        
    md += "## ⚙️ Technologies & Tools Stack\n"
    if report.get("technologies"):
        md += "### Technologies:\n"
        for t in report["technologies"]:
            md += f"- **{t.get('name')}** ({t.get('category')}) - *Level*: {t.get('learning_level')} | *Usage*: {t.get('usage_context')} | *Mentions*: {t.get('mentioned_count', 1)}\n"
            
    if report.get("frameworks"):
        md += "\n### Frameworks:\n"
        for f in report["frameworks"]:
            md += f"- **{f.get('name')}** - *Purpose*: {f.get('purpose')} | *Usage*: {f.get('industry_usage')} | *Resume Worthiness*: {f.get('resume_worthiness')}/100\n"

    if report.get("tools"):
        md += "\n### Tools & Software:\n"
        for t in report["tools"]:
            md += f"- **{t.get('name')}** ({t.get('category')}) - *Context*: {t.get('usage_context')} | *Level*: {t.get('skill_level')}\n"

    md += "\n## 📄 ATS Keyword Optimization\n"
    for k in report.get("ats_keywords", []):
        md += f"- **{k.get('keyword')}** (Score Impact: +{k.get('score_contribution')} | Industry Demand: {k.get('demand_score')}/100)\n"
        
    md += "\n## 📝 High-Impact Resume Achievements\n"
    for bullet in report.get("resume_bullets", []):
        md += f"- {bullet}\n"
        
    md += "\n## 🌐 LinkedIn Profile Boosters\n"
    ls = report.get("linkedin_suggestions", {})
    if ls.get("skills"):
        md += f"- **Add Skills**: {', '.join(ls.get('skills'))}\n"
    if ls.get("headline_suggestions"):
        md += "- **Suggested Headlines**:\n"
        for h in ls["headline_suggestions"]:
            md += f"  - *{h}*\n"
    if ls.get("about_additions"):
        md += f"- **Summary Update**: {ls.get('about_additions')}\n"
    if ls.get("suggested_certifications"):
        md += f"- **Target Certifications**: {', '.join(ls.get('suggested_certifications'))}\n"

    md += "\n## 🎙️ Interview Preparation Q&A\n"
    for idx, q in enumerate(report.get("interview_preparation", [])):
        md += f"### Q{idx+1}: {q.get('question')}\n"
        md += f"**Answer**: {q.get('answer')}\n"
        md += f"*Difficulty*: {q.get('difficulty')} • *Category*: Concept\n\n"
        
    return md

class CareerPDF(FPDF):
    def __init__(self, title_text):
        super().__init__()
        self.title_text = title_text
        
    def header(self):
        if self.page_no() > 1:
            self.set_font('Helvetica', 'I', 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, self.title_text, 0, 0, 'L')
            self.line(10, 17, 200, 17)
            self.ln(10)
            
    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def create_career_pdf(report, title):
    pdf = CareerPDF(f"AI Career Report: {title}")
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Title Cover Page
    pdf.add_page()
    pdf.set_text_color(75, 0, 130) # Indigo
    pdf.set_font("Helvetica", 'B', 22)
    pdf.ln(60)
    pdf.cell(0, 15, "AI Career Intelligence Report", ln=True, align='C')
    pdf.set_font("Helvetica", 'I', 16)
    title_clean = title.encode('latin-1', 'replace').decode('latin-1')
    pdf.cell(0, 10, title_clean, ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Helvetica", size=14)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, "Generated by AI Career Intelligence Suite", ln=True, align='C')
    pdf.cell(0, 10, datetime.now().strftime('%B %d, %Y'), ln=True, align='C')
    
    # 1. Skills Section
    pdf.add_page()
    pdf.set_text_color(75, 0, 130)
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "1. TECHNICAL SKILLS EXTRACTED", ln=True)
    pdf.ln(5)
    
    for s in report.get("skills_learned", []):
        pdf.set_font("Helvetica", 'B', 12)
        pdf.set_text_color(106, 90, 205)
        s_name = s.get('name', 'Skill').encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 8, s_name, ln=True)
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(30, 30, 30)
        desc = f"Description: {s.get('description', '')}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, desc)
        meta = f"Importance: {s.get('importance_score')}/100 | Industry Demand: {s.get('demand_score')}/100 | Resume Relevance: {s.get('resume_relevance_score')}/100".encode('latin-1', 'replace').decode('latin-1')
        pdf.set_font("Helvetica", 'I', 8)
        pdf.set_text_color(128, 128, 128)
        pdf.cell(0, 6, meta, ln=True)
        pdf.ln(4)
        
    # 2. Technologies & Tools
    pdf.add_page()
    pdf.set_text_color(75, 0, 130)
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "2. TECHNOLOGIES, FRAMEWORKS & TOOLS", ln=True)
    pdf.ln(5)
    
    pdf.set_text_color(30, 30, 30)
    if report.get("technologies"):
        pdf.set_font("Helvetica", 'B', 12)
        pdf.cell(0, 8, "Technologies:", ln=True)
        pdf.set_font("Helvetica", size=10)
        for t in report["technologies"]:
            t_str = f"- {t.get('name')} ({t.get('category')}) - Level: {t.get('learning_level')} | Mentions: {t.get('mentioned_count', 1)}".encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, t_str)
        pdf.ln(4)
        
    if report.get("frameworks"):
        pdf.set_font("Helvetica", 'B', 12)
        pdf.cell(0, 8, "Frameworks:", ln=True)
        pdf.set_font("Helvetica", size=10)
        for f in report["frameworks"]:
            f_str = f"- {f.get('name')} - Purpose: {f.get('purpose')} | Resume Worthiness: {f.get('resume_worthiness')}/100".encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, f_str)
        pdf.ln(4)
        
    if report.get("tools"):
        pdf.set_font("Helvetica", 'B', 12)
        pdf.cell(0, 8, "Tools:", ln=True)
        pdf.set_font("Helvetica", size=10)
        for t in report["tools"]:
            t_str = f"- {t.get('name')} ({t.get('category')}) - Skill Level: {t.get('skill_level')}".encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, t_str)
        pdf.ln(4)

    # 3. Resume Bullets & ATS keywords
    pdf.add_page()
    pdf.set_text_color(75, 0, 130)
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "3. RESUME ACHIVEMENTS & ATS OPTIMIZATION", ln=True)
    pdf.ln(5)
    
    pdf.set_font("Helvetica", 'B', 12)
    pdf.cell(0, 8, "ATS-Friendly Optimization Keywords:", ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.set_text_color(30, 30, 30)
    for k in report.get("ats_keywords", []):
        k_str = f"- {k.get('keyword')} (Weight: +{k.get('score_contribution')} | Industry Demand: {k.get('demand_score')}/100)".encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 6, k_str, ln=True)
    pdf.ln(6)
    
    pdf.set_font("Helvetica", 'B', 12)
    pdf.set_text_color(75, 0, 130)
    pdf.cell(0, 8, "ATS Resume Achievements (Bullet Points):", ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.set_text_color(30, 30, 30)
    for bullet in report.get("resume_bullets", []):
        bullet_clean = f"- {bullet}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, bullet_clean)
        pdf.ln(2)
        
    return pdf.output(dest='S')

def create_career_docx(report, title):
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("AI Career Intelligence & Resume Report")
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_paragraph()
    
    # 1. Skills
    h = doc.add_paragraph()
    run = h.add_run("1. Technical Skills Extracted")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    for s in report.get("skills_learned", []):
        p_title = doc.add_paragraph()
        run = p_title.add_run(s.get('name', 'Skill'))
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(106, 90, 205)
        
        p_body = doc.add_paragraph()
        run_b = p_body.add_run(f"Description: {s.get('description', '')}\n")
        run_b.font.name = 'Arial'
        run_b.font.size = Pt(10)
        run_m = p_body.add_run(f"Importance: {s.get('importance_score')}/100 | Demand: {s.get('demand_score')}/100 | Relevance: {s.get('resume_relevance_score')}/100")
        run_m.font.italic = True
        run_m.font.size = Pt(9)
        
    doc.add_paragraph()
    
    # 2. Technologies
    h = doc.add_paragraph()
    run = h.add_run("2. Technologies & Tools Stack")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    if report.get("technologies"):
        p = doc.add_paragraph()
        run_h = p.add_run("Technologies:\n")
        run_h.bold = True
        for t in report["technologies"]:
            p.add_run(f"- {t.get('name')} ({t.get('category')}) - Level: {t.get('learning_level')} | Mentions: {t.get('mentioned_count', 1)}\n")
            
    if report.get("frameworks"):
        p = doc.add_paragraph()
        run_h = p.add_run("Frameworks:\n")
        run_h.bold = True
        for f in report["frameworks"]:
            p.add_run(f"- {f.get('name')} - Purpose: {f.get('purpose')} | Resume Score: {f.get('resume_worthiness')}/100\n")

    if report.get("tools"):
        p = doc.add_paragraph()
        run_h = p.add_run("Tools & Software:\n")
        run_h.bold = True
        for t in report["tools"]:
            p.add_run(f"- {t.get('name')} ({t.get('category')}) - Level: {t.get('skill_level')}\n")

    doc.add_paragraph()
    
    # 3. Resume optimization
    h = doc.add_paragraph()
    run = h.add_run("3. Resume Optimization & ATS Keywords")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    p = doc.add_paragraph()
    run_h = p.add_run("ATS Optimization Keywords:\n")
    run_h.bold = True
    for k in report.get("ats_keywords", []):
        p.add_run(f"- {k.get('keyword')} (Impact: +{k.get('score_contribution')} | Demand: {k.get('demand_score')}/100)\n")
        
    doc.add_paragraph()
    p = doc.add_paragraph()
    run_h = p.add_run("ATS Achievement Bullets:\n")
    run_h.bold = True
    for bullet in report.get("resume_bullets", []):
        p.add_run(f"- {bullet}\n")
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
