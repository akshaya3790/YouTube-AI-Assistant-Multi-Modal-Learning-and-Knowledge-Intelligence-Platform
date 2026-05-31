import os
# Fix for OpenMP duplicate library error - MUST be at the very top
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

import streamlit as st
import yt_dlp
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage
import time
import json
from datetime import datetime
from fpdf import FPDF
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from youtube_transcript_api import YouTubeTranscriptApi
import zipfile
import re
import io
import rag_storage
from rag_storage import VideoRAGStore
import graph_generator
import study_storage
from study_storage import StudyStorage
import study_generator
import document_builder
import export_manager
import multi_video_generator
from multi_video_storage import MultiVideoStorage

# Configure Gemini
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY") or "AIzaSyCefTl5p31gixkgkTFfaO7WJefK_VbB_RQ"
genai.configure(api_key=GEMINI_API_KEY)
os.environ["GOOGLE_API_KEY"] = GEMINI_API_KEY

st.set_page_config(
    layout="wide",
    page_title="YouTube AI Assistant",
    page_icon="🎬"
)

# Initialize Session States
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'detailed_summary' not in st.session_state:
    st.session_state.detailed_summary = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'video_info' not in st.session_state:
    st.session_state.video_info = None
if 'processed_url' not in st.session_state:
    st.session_state.processed_url = None
if 'quiz_data' not in st.session_state:
    st.session_state.quiz_data = None
if 'quiz_submitted' not in st.session_state:
    st.session_state.quiz_submitted = False
if 'user_answers' not in st.session_state:
    st.session_state.user_answers = {}
if 'view' not in st.session_state:
    st.session_state.view = 'main'
if 'selected_model' not in st.session_state:
    st.session_state.selected_model = "gemini-pro-latest"

def download_audio(url):
    try:
        if not os.path.exists("temp_audio"):
            os.makedirs("temp_audio")
            
        ydl_opts = {
            'format': 'm4a/bestaudio/best',
            'outtmpl': 'temp_audio/%(id)s.%(ext)s',
            'quiet': True,
            'no_warnings': True,
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            filename = ydl.prepare_filename(info)
            return filename, info
    except Exception as e:
        st.error(f"Error downloading video: {e}")
        return None, None

def create_pdf_bytes(content, title):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", 'B', 16)
        pdf.cell(0, 10, title, ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Helvetica", size=12)
        clean_content = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, clean_content)
        return pdf.output()
    except Exception as e:
        st.error(f"Error creating PDF: {e}")
        return None

def process_with_gemini(file_path, start_time=None, end_time=None, summary_type="Executive Summary"):
    try:
        duration = st.session_state.video_info.get('duration', 0)
        is_long_video = duration > 3600
        
        with st.spinner(f"📤 Uploading to AI (this might take a moment for long videos)..."):
            audio_file = genai.upload_file(path=file_path)
        
        with st.spinner(f"🧠 {st.session_state.selected_model} is analyzing the content (this may take several minutes for videos over 1 hour)..."):
            wait_start_time = time.time()
            while audio_file.state.name == "PROCESSING":
                time.sleep(5) # Check less frequently for long videos
                audio_file = genai.get_file(audio_file.name)
                
                # 10 minute timeout
                if time.time() - wait_start_time > 600:
                    raise Exception("Processing timed out after 10 minutes. The server might be overloaded.")
            
        if audio_file.state.name == "FAILED":
            raise Exception("AI failed to process the audio file. It might be corrupted or unsupported.")
            
        model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
        
        time_constraint = ""
        if start_time or end_time:
            time_constraint = f"\n\nIMPORTANT: ONLY analyze and summarize the specific segment of the video "
            if start_time:
                time_constraint += f"starting from {start_time} "
            if end_time:
                if not start_time:
                    time_constraint += "from the beginning "
                time_constraint += f"up to {end_time}."
            time_constraint += "\nDo not include information outside of this timeframe."

        # Define prompts based on selected type
        if summary_type == "Tweet-Length (Very Short)":
            prompt = f"""
            Analyze this video audio and provide a response in Markdown:{time_constraint}
            
            Write a highly engaging, tweet-length summary (under 280 characters) of the main point of this video. Include 2-3 relevant hashtags at the end.
            """
        elif summary_type == "Bullet Points (Quick Read)":
            prompt = f"""
            Analyze this video audio and provide a response in Markdown:{time_constraint}
            
            Provide a concise, bullet-point only summary of the most important facts, arguments, and takeaways from this video. Do not use paragraphs, just clear, crisp bullet points.
            """
        elif summary_type == "Simple Explanation (Like I'm 5)":
            prompt = f"""
            Analyze this video audio and provide a response in Markdown:{time_constraint}
            
            Explain the core concepts of this video in extremely simple terms, as if you were explaining it to a 5-year-old. Use analogies and avoid any complex jargon.
            """
        elif summary_type == "Detailed Report (Comprehensive)":
            prompt = f"""
            Analyze this video audio and provide a response in Markdown:{time_constraint}
            
            Please act as an expert analyst and provide an EXTREMELY DETAILED, comprehensive report.
            Include:
            ## 📝 Executive Summary
            ## 📘 Detailed Chapter Breakdown (chronological/logical flow)
            ## 💡 Key Takeaways & Actionable Insights
            ## 🌐 Deep Dive & Related Resources
            """
        else: # Standard "Executive Summary"
            prompt = f"""
            Analyze this video audio and provide a response in Markdown:{time_constraint}
            
            ## 📝 Executive Summary
            (3-4 paragraphs summarizing the main content)
            
            ## 💡 Key Takeaways
            (5-7 bullet points of main points)
            
            ## 🌐 Deep Dive & Related Resources
            (Suggest 3-4 topics and 2-3 websites/channels)
            
            ## 🎭 Tone & Sentiment
            (Brief vibe check)
            """
        
        # Retry logic for 429 errors
        max_api_retries = 3
        for attempt in range(max_api_retries):
            try:
                message = HumanMessage(
                    content=[
                        {"type": "text", "text": prompt},
                        {"type": "media", "file_uri": audio_file.uri, "mime_type": audio_file.mime_type}
                    ]
                )
                response = model.invoke([message])
                break
            except Exception as e:
                if "429" in str(e) and attempt < max_api_retries - 1:
                    st.warning(f"Quota exceeded for {st.session_state.selected_model}. Retrying in 10s... (Attempt {attempt + 1})")
                    time.sleep(10)
                else:
                    raise e
                    
        st.session_state.audio_file_name = audio_file.name
        return response.content
    except Exception as e:
        if "429" in str(e):
            st.error("🚨 Quota Exceeded! Please switch to a different AI model in the sidebar and try again.")
        else:
            st.error(f"Error processing with AI: {e}")
        return None

def generate_quiz():
    try:
        duration = st.session_state.video_info.get('duration', 0)
        num_questions = 10 if duration > 1800 else 5
        
        model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
        chat_context = genai.get_file(st.session_state.audio_file_name)
        prompt = f"""
        Based on this video, generate a {num_questions}-question Multiple Choice Quiz.
        Return ONLY a JSON array with question, options, and answer keys.
        """
        message = HumanMessage(
            content=[
                {"type": "text", "text": prompt},
                {"type": "media", "file_uri": chat_context.uri, "mime_type": chat_context.mime_type}
            ]
        )
        response = model.invoke([message])
        json_text = response.content.replace('```json', '').replace('```', '').strip()
        return json.loads(json_text)
    except Exception as e:
        st.error(f"Error generating quiz: {e}")
        return None

class BeautifulPDF(FPDF):
    def __init__(self, title_text):
        super().__init__()
        self.title_text = title_text
        
    def header(self):
        if self.page_no() > 1:
            self.set_font('Helvetica', 'I', 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, self.title_text, 0, 0, 'L')
            self.line(10, 17, 200, 17)
            self.ln(10)
            
    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def create_beautiful_pdf_bytes(content, title):
    pdf = BeautifulPDF(title)
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # 1. Cover Page
    pdf.add_page()
    pdf.set_text_color(75, 0, 130) # Indigo
    pdf.set_font("Helvetica", 'B', 24)
    pdf.ln(60)
    pdf.cell(0, 15, title, ln=True, align='C')
    pdf.set_text_color(100, 100, 100)
    pdf.set_font("Helvetica", size=14)
    pdf.ln(10)
    pdf.cell(0, 10, "Generated by AI Notes Assistant", ln=True, align='C')
    pdf.cell(0, 10, datetime.now().strftime('%B %d, %Y'), ln=True, align='C')
    
    # 2. Main content page
    pdf.add_page()
    pdf.set_text_color(0, 0, 0)
    
    for line in content.split('\n'):
        # Clean latin-1 encoding issues
        line = line.encode('latin-1', 'replace').decode('latin-1')
        
        if line.startswith('# '):
            pdf.ln(5)
            pdf.set_font("Helvetica", 'B', 18)
            pdf.set_text_color(75, 0, 130)
            pdf.cell(0, 10, line[2:], ln=True)
            pdf.ln(3)
        elif line.startswith('## '):
            pdf.ln(4)
            pdf.set_font("Helvetica", 'B', 14)
            pdf.set_text_color(106, 90, 205)
            pdf.cell(0, 8, line[3:], ln=True)
            pdf.ln(2)
        elif line.startswith('### '):
            pdf.ln(3)
            pdf.set_font("Helvetica", 'B', 12)
            pdf.set_text_color(50, 50, 50)
            pdf.cell(0, 6, line[4:], ln=True)
            pdf.ln(1)
        else:
            pdf.set_font("Helvetica", size=10)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 6, line)
            pdf.ln(1)
            
    return pdf.output(dest='S')

def create_docx_bytes(content, title):
    doc = Document()
    
    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130) # Indigo
    
    doc.add_paragraph()
    
    # Content
    for line in content.split('\n'):
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(75, 0, 130)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(106, 90, 205) # SlateBlue
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_zip_package(notes_dict, title_safe):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for note_type, content in notes_dict.items():
            safe_type_name = note_type.replace(" ", "_")
            # 1. MD
            zip_file.writestr(f"{title_safe}_{safe_type_name}.md", content)
            # 2. PDF
            pdf_bytes = create_beautiful_pdf_bytes(content, f"{note_type}: {title_safe}")
            zip_file.writestr(f"{title_safe}_{safe_type_name}.pdf", pdf_bytes)
            # 3. DOCX
            docx_bytes = create_docx_bytes(content, f"{note_type}: {title_safe}")
            zip_file.writestr(f"{title_safe}_{safe_type_name}.docx", docx_bytes)
            
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def get_youtube_transcript(url):
    if url.startswith("local://"):
        from media_library import MediaLibrary
        media_id = url.split("local://")[1]
        lib = MediaLibrary()
        media = lib.get_media(media_id)
        if media and media.get("transcript"):
            return media["transcript"]
        return None

    video_id = None
    patterns = [
        r"(?:v=|\/v\/|\/embed\/|\/shorts\/|youtu\.be\/)([a-zA-Z0-9_-]{11})"
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            video_id = match.group(1)
            break
            
    if not video_id:
        return None
        
    try:
        print("VIDEO ID:", video_id)
        print("TRANSCRIPT API START")
        
        # Version-safe API retrieval (Handles v1.2+ and older versions)
        try:
            # Try new instantiated API format (v1.2+)
            api_instance = YouTubeTranscriptApi()
            transcript_list_obj = api_instance.list(video_id)
            transcript_obj = transcript_list_obj.find_transcript(['en', 'hi', 'te', 'ta', 'kn'])
            raw_transcript = transcript_obj.fetch()
            
            transcript_list = []
            for item in raw_transcript:
                text = getattr(item, 'text', item.get('text', ''))
                transcript_list.append({'text': text})
        except AttributeError:
            # Fallback for older versions
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en', 'hi', 'te', 'ta', 'kn'])
            
        print("TRANSCRIPT RETRIEVED")
        transcript_text = " ".join([item['text'] for item in transcript_list])
        return transcript_text
    except Exception as e:
        print("TRANSCRIPT FAILED:", str(e))
        print("WHISPER FALLBACK STARTED")
        st.warning(f"Using Whisper Fallback (Transcript Unavailable: {e})")
        return None

def transcribe_media_with_gemini(file_path):
    try:
        with st.spinner("📤 Uploading media to AI for transcription..."):
            audio_file = genai.upload_file(path=file_path)
            
        with st.spinner("🧠 Transcribing audio/video content (this may take a few minutes for long media)..."):
            wait_start_time = time.time()
            while audio_file.state.name == "PROCESSING":
                time.sleep(5)
                audio_file = genai.get_file(audio_file.name)
                if time.time() - wait_start_time > 600:
                    raise Exception("Transcription processing timed out.")
                    
        if audio_file.state.name == "FAILED":
            raise Exception("AI failed to process the audio file.")
            
        model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
        prompt = "Listen to the uploaded audio/video and generate a full, word-for-word transcript. Remove filler words like um, uh, you know, etc. and output only the clean transcript."
        
        message = HumanMessage(
            content=[
                {"type": "text", "text": prompt},
                {"type": "media", "file_uri": audio_file.uri, "mime_type": audio_file.mime_type}
            ]
        )
        response = model.invoke([message])
        return response.content
    except Exception as e:
        st.error(f"Error transcribing media: {e}")
        return None

def detect_category(transcript_snippet):
    try:
        model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
        prompt = f"""
        Analyze the following snippet of a video transcript and classify it into exactly one of these categories:
        Education, Programming, Data Science, AI/ML, Business, Finance, Interview, Webinar, Meeting, Podcast, Research, Healthcare, Technology, General Knowledge.
        
        Return ONLY the category name.
        
        Snippet:
        {transcript_snippet[:2000]}
        """
        response = model.invoke([HumanMessage(content=prompt)])
        category = response.content.strip()
        return category
    except Exception as e:
        return "General Knowledge"

def extract_metadata(transcript):
    try:
        model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
        prompt = f"""
        Analyze the following transcript and extract key entities and metadata.
        Provide the output in valid JSON format with the following keys:
        - important_concepts (list of strings)
        - definitions (dict where keys are terms and values are definitions)
        - technical_terms (list of strings)
        - formulas (list of strings)
        - frameworks (list of strings)
        - algorithms (list of strings)
        - tools_libraries_software (list of strings)
        - company_names (list of strings)
        - research_references (list of strings)
        - statistics (list of strings)
        - action_items (list of strings)
        - recommendations (list of strings)

        Return ONLY the raw JSON. No markdown code blocks, no other text.
        
        Transcript:
        {transcript[:8000]}
        """
        response = model.invoke([HumanMessage(content=prompt)])
        json_text = response.content.replace('```json', '').replace('```', '').strip()
        return json.loads(json_text)
    except Exception as e:
        return {}

def parse_ts_to_secs(ts_str):
    try:
        parts = list(map(int, ts_str.split(':')))
        if len(parts) == 2:
            return parts[0] * 60 + parts[1]
        elif len(parts) == 3:
            return parts[0] * 3600 + parts[1] * 60 + parts[2]
    except:
        pass
    return 0.0

def format_secs_to_ts(seconds):
    s = int(seconds)
    h = s // 3600
    m = (s % 3600) // 60
    sec = s % 60
    if h > 0:
        return f"{h:02d}:{m:02d}:{sec:02d}"
    return f"{m:02d}:{sec:02d}"

def find_chapter_for_time(seconds, chapters):
    if not chapters:
        return "General Content"
    current_chapter = "Introduction"
    for ch in chapters:
        start_s = ch.get('seconds')
        if start_s is None and 'start_time' in ch:
            start_s = parse_ts_to_secs(ch['start_time'])
        
        if start_s is not None and seconds >= start_s:
            current_chapter = ch.get('title', 'Chapter')
        elif start_s is not None:
            break
    return current_chapter

def preprocess_transcript(raw_text):
    if not raw_text:
        return ""
    fillers = [r'\buh\b', r'\bum\b', r'\blike\b', r'\bso\b', r'\bactually\b', r'\bbasically\b', r'\byou know\b']
    cleaned = raw_text
    for filler in fillers:
        cleaned = re.sub(filler, '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned

def create_mock_transcript_list(text, duration_seconds=600):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    total_chars = sum(len(s) for s in sentences)
    if total_chars == 0:
        return []
        
    transcript_list = []
    current_sec = 0.0
    for s in sentences:
        if not s.strip():
            continue
        frac = len(s) / total_chars
        dur = frac * duration_seconds
        transcript_list.append({
            "text": s.strip(),
            "start": current_sec,
            "duration": dur
        })
        current_sec += dur
    return transcript_list

def chunk_transcript(raw_transcript_list, chapters, chunk_size_words=600, overlap_pct=0.15):
    chunks = []
    metadatas = []
    
    if not raw_transcript_list:
        return chunks, metadatas
        
    i = 0
    n = len(raw_transcript_list)
    overlap_words = int(chunk_size_words * overlap_pct)
    
    current_segments = []
    current_word_count = 0
    
    while i < n:
        seg = raw_transcript_list[i]
        words = seg['text'].split()
        current_segments.append(seg)
        current_word_count += len(words)
        i += 1
        
        if current_word_count >= chunk_size_words or i == n:
            chunk_text = " ".join([s['text'] for s in current_segments])
            start_secs = current_segments[0]['start']
            end_secs = current_segments[-1]['start'] + current_segments[-1].get('duration', 0)
            
            start_ts = format_secs_to_ts(start_secs)
            end_ts = format_secs_to_ts(end_secs)
            
            ch_name = find_chapter_for_time(start_secs, chapters)
            
            chunks.append(chunk_text)
            metadatas.append({
                "start_time": start_ts,
                "end_time": end_ts,
                "start_seconds": start_secs,
                "end_seconds": end_secs,
                "chapter_name": ch_name,
                "topic": "General",
                "keywords": ""
            })
            
            overlap_segs = []
            overlap_count = 0
            for j in range(len(current_segments)-1, -1, -1):
                seg_words = len(current_segments[j]['text'].split())
                if overlap_count + seg_words <= overlap_words:
                    overlap_segs.insert(0, current_segments[j])
                    overlap_count += seg_words
                else:
                    break
            
            current_segments = overlap_segs
            current_word_count = overlap_count
            
    return chunks, metadatas

def generate_embedding(text):
    try:
        result = genai.embed_content(
            model="models/gemini-embedding-001",
            content=text,
            task_type="retrieval_document"
        )
        return result["embedding"]
    except Exception as e:
        print(f"Embedding error: {e}")
        return [0.0] * 3072

def generate_embeddings_batch(texts):
    try:
        result = genai.embed_content(
            model="models/gemini-embedding-001",
            content=texts,
            task_type="retrieval_document"
        )
        return result["embedding"]
    except Exception as e:
        print(f"Batch embedding error: {e}")
        embeddings = []
        for text in texts:
            embeddings.append(generate_embedding(text))
        return embeddings

def generate_rag_answer(query, history, contexts, mode, selected_model="gemini-2.5-flash"):
    try:
        context_str = ""
        for i, ctx in enumerate(contexts):
            meta = ctx["metadata"]
            context_str += f"[Doc {i+1}] (Chapter: {meta.get('chapter_name', 'N/A')}, Timestamp: {meta.get('start_time', '00:00')} - {meta.get('end_time', '00:00')})\nContent: {ctx['text']}\n\n"
            
        history_str = ""
        for h in history[-5:]:
            role = "User" if h["role"] == "user" else "AI Assistant"
            history_str += f"{role}: {h['content']}\n"
            
        mode_instruction = ""
        if mode == "Explain Simply":
            mode_instruction = "Explain in plain, simple English suitable for a non-technical person. Use analogies."
        elif mode == "Explain Like I'm 10":
            mode_instruction = "Explain like I'm 10 years old. Use highly engaging simple analogies and short sentences."
        elif mode == "Explain Technically":
            mode_instruction = "Explain using advanced technical terms, architectural designs, algorithms, code libraries, and technical workflows."
        elif mode == "Interview Mode":
            mode_instruction = "Format your answer as a QA dialogue suitable for a technical interview prep, highlighting candidate responses, follow-up questions, and key keywords to mention."
        elif mode == "Exam Mode":
            mode_instruction = "Format your answer with clear definitions, bullet points of key takeaways, formulas if any, and list of exam-oriented remember items."
        elif mode == "Developer Mode":
            mode_instruction = "Focus on software engineering practices, code structures, libraries, APIs, database architectures, and implementation patterns."
        elif mode == "Research Mode":
            mode_instruction = "Focus on details, analytical breakdowns, relationships, citations, and methodologies."
            
        prompt = f"""
        You are a Senior AI Learning Assistant. Your goal is to answer the user's question based strictly on the provided video transcript contexts.
        
        Guidelines:
        1. Stay grounded in the context. If the answer cannot be found in the provided context, state that clearly and do not hallucinate.
        2. Format your response based on the selected assistant mode.
        3. Do not refer to "documents" or "provided text" directly in your response text (e.g. say "The speaker explains" instead of "According to Document 1").
        
        Active Assistant Mode: {mode}
        Mode Instructions: {mode_instruction}
        
        Transcript Contexts:
        {context_str}
        
        Conversation History:
        {history_str}
        
        User Question: {query}
        
        Provide your formatted answer:
        """
        
        model = ChatGoogleGenerativeAI(model=selected_model)
        response = model.invoke([HumanMessage(content=prompt)])
        return response.content
    except Exception as e:
        return f"Error generating answer: {e}"

def index_video_in_rag(url, title, info, transcript_text=None, transcript_list=None):
    try:
        store = VideoRAGStore(path="chroma_db_storage")
        
        video_id = None
        patterns = [r"(?:v=|\/v\/|\/embed\/|\/shorts\/|youtu\.be\/)([a-zA-Z0-9_-]{11})"]
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                video_id = match.group(1)
                break
        if not video_id:
            video_id = "upload_" + str(hash(title) & 0xffffffff)
            
        if not transcript_text or not transcript_list:
            if "youtube.com" in url or "youtu.be" in url:
                transcript_text, transcript_list = get_youtube_transcript_with_timestamps(url)
            
        if not transcript_text:
            return False
            
        cleaned_text = preprocess_transcript(transcript_text)
        
        if not transcript_list:
            duration = info.get("duration", 600) if info else 600
            transcript_list = create_mock_transcript_list(cleaned_text, duration)
            
        original_chapters = get_youtube_original_chapters(info) if info else []
        
        chunks, metadatas = chunk_transcript(transcript_list, original_chapters)
        
        if not chunks:
            return False
            
        embeddings = generate_embeddings_batch(chunks)
        store.add_transcript_chunks(video_id, title, chunks, embeddings, metadatas)
        return True
    except Exception as e:
        print(f"Error indexing video in RAG: {e}")
        return False

def get_category_guidelines(category):
    guidelines = ""
    if category == "Programming":
        guidelines = "Prioritize: code concepts, libraries, frameworks, syntax explanations, and implementation details."
    elif category in ["Business", "Finance", "Meeting"]:
        guidelines = "Prioritize: decisions, business strategies, action plans, KPIs, and organizational goals."
    elif category in ["Education", "Research", "AI/ML", "Data Science"]:
        guidelines = "Prioritize: core definitions, scientific concepts, mathematical formulas, algorithms, academic frameworks, and detailed examples."
    else:
        guidelines = "Prioritize: core themes, important facts, definitions, and key takeaways."
    return guidelines

def generate_notes_chunk(chunk_text, note_type, category):
    model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
    category_guidelines = get_category_guidelines(category)
    
    if note_type == "Detailed Notes":
        prompt = f"""
        You are an expert learning assistant. Generate comprehensive, highly detailed learning notes from the following transcript.
        Follow these requirements:
        - Maintain the logical flow of the original content.
        - Include all important explanations and preserve technical concepts.
        - Include examples mentioned in the transcript.
        - Explain complex terms in detail.
        - Organize the information hierarchically with proper headings and subheadings.
        - Avoid unnecessary filler content.
        - Category-specific focus: {category_guidelines}

        Use this exact structure:
        # Detailed Notes
        ## 📋 Overview
        [Overview of the content]

        ## 🔍 Main Topics
        ### [Topic 1 Name]
        - **Explanation**: [Detailed explanation]
        - **Key Concepts**: [Key concepts explained]
        - **Examples**: [Examples mentioned]

        ## 💡 Important Observations
        [Key insights, nuances, or observations]

        ## 🏁 Conclusion
        [Final summary and takeaways]

        Transcript:
        {chunk_text}
        """
    elif note_type == "Study Notes":
        prompt = f"""
        You are an expert educator. Generate exam-oriented study notes from the following transcript.
        Follow these requirements:
        - Focus on key concepts and definitions.
        - Highlight important formulas, frameworks, and technical keywords.
        - Emphasize interview-relevant points.
        - Category-specific focus: {category_guidelines}

        Use this exact structure:
        # Study Notes
        ## 📘 Topic Overview
        [Summary of the topic]

        ## 📕 Important Definitions
        - **[Term]**: [Definition]

        ## 🧠 Key Concepts
        [Detailed explanation of the key concepts]

        ## 📊 Important Facts
        [Facts, statistics, and keywords to remember]

        ## ➗ Formulas / Frameworks
        [Any mathematical formulas, software architectures, or theoretical frameworks]

        ## 💡 Examples
        [Practical examples mentioned]

        ## ❓ Frequently Asked Questions
        [Generate FAQs based on the content with answers]

        ## ⚡ Quick Revision Section
        [A bullet-point summary for quick scanning]

        Transcript:
        {chunk_text}
        """
    elif note_type == "Meeting Notes":
        prompt = f"""
        You are a professional scribe. Convert the following transcript into high-quality corporate meeting minutes.
        Follow these requirements:
        - Automatically identify the meeting purpose, key discussions, and recommendations.
        - List all key decisions, action items, responsibilities, and deadlines.
        - Category-specific focus: {category_guidelines}

        Use this exact structure:
        # Meeting Minutes
        ## 📝 Meeting Summary
        [Executive summary of the discussion]

        ## 🎯 Objectives
        [Objectives of the meeting]

        ## 💬 Discussion Topics
        ### [Topic]
        [Summary of what was discussed]

        ## ⚖️ Key Decisions
        [Decisions made during the meeting]

        ## 📋 Action Items
        - **[Action Item]**: [Details] (Assigned to: [Person], Deadline: [Date/Time])

        ## 👥 Responsible Persons & Deadlines
        [Summary of responsibilities and deadlines]

        ## 🚀 Next Steps
        [Next steps for the team]

        ## 🏁 Conclusion
        [Closing remarks]

        Transcript:
        {chunk_text}
        """
    else:  # Revision Notes
        prompt = f"""
        You are an expert summarizer. Generate highly compressed revision notes from the following transcript.
        Follow these requirements:
        - Extract only critical information and remove all unnecessary explanations.
        - Focus on memorization and quick learning.
        - Category-specific focus: {category_guidelines}

        Use this exact structure:
        # Revision Notes
        ## ⚡ Core Concepts
        - [Key concept]

        ## 🔑 Important Keywords
        [Comma-separated list of key terms]

        ## 📖 Definitions
        - **[Term]**: [Definition]

        ## 📌 Facts to Remember
        [Key statistics, facts, or bullet points]

        ## 📝 Quick Bullet Points
        [Rapid-fire summaries]

        ## ⏱️ One-Minute Revision
        [A single paragraph summarizing the absolute core message]

        Transcript:
        {chunk_text}
        """
        
    response = model.invoke([HumanMessage(content=prompt)])
    return response.content

def merge_notes(chunk_notes, note_type, title):
    model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
    prompt = f"""
    You are an expert editor. You have been given notes generated from different parts of a video transcript.
    Your task is to merge them into a single, cohesive, professionally formatted set of notes for the note type: '{note_type}'.
    
    Ensure that:
    - The final notes are well-organized and avoid duplicate headings or redundant information.
    - All technical terms, concepts, examples, and definitions from the individual notes are preserved.
    - The structural template for '{note_type}' is strictly followed.
    - Title should be: '{title}'
    
    Individual notes to merge:
    {"="*40}
    """
    for i, notes in enumerate(chunk_notes):
        prompt += f"\n\n--- Part {i+1} ---\n{notes}"
        
    response = model.invoke([HumanMessage(content=prompt)])
    return response.content

def run_notes_generation(transcript, note_types, title, category):
    words = transcript.split()
    chunk_size = 15000
    chunks = []
    
    for i in range(0, len(words), chunk_size):
        chunk_words = words[i:i+chunk_size]
        chunks.append(" ".join(chunk_words))
        
    generated_notes = {}
    
    for note_type in note_types:
        with st.spinner(f"📝 Generating {note_type}..."):
            chunk_results = []
            for j, chunk in enumerate(chunks):
                if len(chunks) > 1:
                    st.text(f"Processing chunk {j+1} of {len(chunks)} for {note_type}...")
                chunk_res = generate_notes_chunk(chunk, note_type, category)
                chunk_results.append(chunk_res)
                
            if len(chunk_results) == 1:
                final_note = f"# {title}\n\n" + chunk_results[0]
            else:
                st.text(f"Merging chunk notes for {note_type}...")
                final_note = merge_notes(chunk_results, note_type, title)
                
            generated_notes[note_type] = final_note
            
    return generated_notes

def notes_generator_view():
    st.header("📝 AI Notes Generator")
    st.write("Transform video transcripts and audio/video files into beautifully structured learning materials, study guides, corporate meeting minutes, or quick revision sheets.")
    
    if 'notes_transcript' not in st.session_state:
        st.session_state.notes_transcript = None
    if 'notes_title' not in st.session_state:
        st.session_state.notes_title = "Uploaded Media Notes"
    if 'notes_category' not in st.session_state:
        st.session_state.notes_category = None
    if 'notes_metadata' not in st.session_state:
        st.session_state.notes_metadata = {}
    if 'notes_results' not in st.session_state:
        st.session_state.notes_results = {}
        
    source_type = st.radio("Choose Input Source", ["YouTube Video URL", "Upload Audio/Video File"], horizontal=True)
    
    if source_type == "YouTube Video URL":
        youtube_url = st.text_input("Enter YouTube Video URL", placeholder="https://www.youtube.com/watch?v=...")
        if st.button("🔍 Get Transcript"):
            if youtube_url:
                st.info("Fetching video info and transcript...")
                audio_file, info = download_audio(youtube_url)
                if info:
                    st.session_state.notes_title = info.get('title', 'YouTube Video')
                
                transcript = get_youtube_transcript(youtube_url)
                if not transcript and audio_file:
                    transcript = transcribe_media_with_gemini(audio_file)
                    
                if audio_file and os.path.exists(audio_file):
                    os.remove(audio_file)
                    
                if transcript:
                    st.session_state.notes_transcript = transcript
                    index_video_in_rag(youtube_url, info.get('title', 'YouTube Video') if info else st.session_state.notes_title, info, transcript)
                    st.success("Transcript generated successfully and indexed for RAG chatbot/mind map!")
                    st.rerun()
                else:
                    st.error("Failed to generate transcript. Please ensure the URL is valid or upload an audio file instead.")
    else:
        uploaded_file = st.file_uploader("Upload Audio or Video File", type=["mp3", "m4a", "wav", "mp4", "mov", "avi", "webm"])
        if uploaded_file is not None:
            st.session_state.notes_title = uploaded_file.name
            if st.button("🎙️ Transcribe Uploaded File"):
                if not os.path.exists("temp_audio"):
                    os.makedirs("temp_audio")
                temp_path = os.path.join("temp_audio", uploaded_file.name)
                with open(temp_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                transcript = transcribe_media_with_gemini(temp_path)
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
                if transcript:
                    st.session_state.notes_transcript = transcript
                    index_video_in_rag(st.session_state.notes_title, st.session_state.notes_title, None, transcript)
                    st.success("File transcribed successfully and indexed for RAG chatbot/mind map!")
                    st.rerun()
                else:
                    st.error("Failed to transcribe file.")

    if st.session_state.notes_transcript:
        st.divider()
        st.subheader("Transcript Preview")
        with st.expander("Show Full Transcript"):
            st.write(st.session_state.notes_transcript)
            
        st.divider()
        st.subheader("Generate Structured Notes")
        
        note_types_options = ["Detailed Notes", "Study Notes", "Meeting Notes", "Revision Notes"]
        selected_types = st.multiselect("Select Note Types to Generate", note_types_options, default=["Detailed Notes", "Study Notes"])
        
        if st.button("🧠 Generate Selected Notes"):
            if selected_types:
                with st.spinner("🔍 Detecting content category..."):
                    st.session_state.notes_category = detect_category(st.session_state.notes_transcript)
                    st.info(f"Category Detected: **{st.session_state.notes_category}**")
                    
                with st.spinner("📊 Extracting key terms, concepts, and definitions..."):
                    st.session_state.notes_metadata = extract_metadata(st.session_state.notes_transcript)
                    
                st.session_state.notes_results = run_notes_generation(
                    st.session_state.notes_transcript,
                    selected_types,
                    st.session_state.notes_title,
                    st.session_state.notes_category
                )
                st.success("All selected notes generated successfully!")
                st.rerun()
            else:
                st.warning("Please select at least one note type to generate.")

    if st.session_state.notes_results:
        st.divider()
        st.subheader("📋 Notes Preview & Export")
        
        tabs = st.tabs(list(st.session_state.notes_results.keys()))
        for idx, (note_type, content) in enumerate(st.session_state.notes_results.items()):
            with tabs[idx]:
                st.markdown(content)
                
                st.divider()
                st.subheader(f"📥 Export options for {note_type}")
                
                c1, c2, c3 = st.columns(3)
                with c1:
                    st.download_button(
                        label="📄 Download Markdown (.md)",
                        data=content,
                        file_name=f"{st.session_state.notes_title.replace(' ', '_')}_{note_type.replace(' ', '_')}.md",
                        mime="text/markdown",
                        key=f"dl_md_{note_type}"
                    )
                with c2:
                    pdf_bytes = create_beautiful_pdf_bytes(content, f"{note_type}: {st.session_state.notes_title}")
                    st.download_button(
                        label="📕 Download PDF",
                        data=pdf_bytes,
                        file_name=f"{st.session_state.notes_title.replace(' ', '_')}_{note_type.replace(' ', '_')}.pdf",
                        mime="application/pdf",
                        key=f"dl_pdf_{note_type}"
                    )
                with c3:
                    docx_bytes = create_docx_bytes(content, f"{note_type}: {st.session_state.notes_title}")
                    st.download_button(
                        label="📘 Download Word (.docx)",
                        data=docx_bytes,
                        file_name=f"{st.session_state.notes_title.replace(' ', '_')}_{note_type.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_docx_{note_type}"
                    )

        st.divider()
        st.subheader("📦 Bulk Export Package")
        zip_bytes = create_zip_package(st.session_state.notes_results, st.session_state.notes_title.replace(' ', '_'))
        st.download_button(
            label="📚 Download all note formats as a ZIP Archive",
            data=zip_bytes,
            file_name=f"{st.session_state.notes_title.replace(' ', '_')}_notes_package.zip",
            mime="application/zip"
        )
        
        if st.session_state.notes_metadata:
            st.divider()
            st.subheader("💡 Smart Extracted Knowledge Metadata")
            
            meta = st.session_state.notes_metadata
            
            if meta.get("important_concepts"):
                st.write("**Important Concepts:**")
                st.write(", ".join(meta["important_concepts"]))
                
            if meta.get("definitions"):
                st.write("**Definitions:**")
                for term, definition in meta["definitions"].items():
                    st.write(f"- **{term}**: {definition}")
                    
            if meta.get("technical_terms"):
                st.write("**Technical Terms & Keywords:**")
                st.write(", ".join(meta["technical_terms"]))
                
            if meta.get("formulas"):
                st.write("**Formulas:**")
                for formula in meta["formulas"]:
                    st.write(f"- `{formula}`")
                    
            if meta.get("frameworks"):
                st.write("**Frameworks & Architectures:**")
                st.write(", ".join(meta["frameworks"]))
                
            if meta.get("tools_libraries_software"):
                st.write("**Tools, Libraries & Software:**")
                st.write(", ".join(meta["tools_libraries_software"]))
                
            if meta.get("action_items"):
                st.write("**Action Items:**")
                for item in meta["action_items"]:
                    st.write(f"- [ ] {item}")
                    
            if meta.get("recommendations"):
                st.write("**Recommendations:**")
                for rec in meta["recommendations"]:
                    st.write(f"- 💡 {rec}")

def generate_intermediate_summary(chunk_text, category):
    model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
    prompt = f"""
    You are an expert analyst. Summarize the following chunk of video transcript.
    Focus on extracting the most important insights, technical terms, decisions, and key points related to the category '{category}'.
    
    Chunk:
    {chunk_text}
    """
    response = model.invoke([HumanMessage(content=prompt)])
    return response.content

def generate_final_takeaways_payload(consolidated_text, title, category):
    model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
    category_guidelines = get_category_guidelines(category)
    
    prompt = f"""
    You are a Senior AI Engineer and NLP Architect. Extract the most important insights and summaries from the following consolidated transcript summaries.
    Provide the output in valid JSON format with the following keys:
    - top_5: a list of exactly 5 dicts, where each dict has:
        - rank: integer (1 to 5)
        - title: string (insight title)
        - explanation: string (2-4 sentences explaining why it matters)
        - practical_value: string (how to apply it)
    - top_10: a list of exactly 10 dicts, where each dict has:
        - rank: integer (1 to 10)
        - title: string (takeaway title)
        - importance_score: integer (1 to 100, based on impact and strategic value)
        - explanation: string (explanation of the insight)
        - practical_application: string (how to apply it)
    - executive_summary: a dict with keys:
        - short: string (100-200 words summary)
        - medium: string (300-500 words summary)
        - detailed: string (600-1000 words summary)
    - analytics: a dict with keys:
        - num_insights_extracted: integer
        - highest_scoring_insight: string (the title of the highest scoring takeaway)
        - most_discussed_topic: string
        - most_impactful_takeaway: string
        - content_complexity_score: integer (1 to 10)
        - learning_value_score: integer (1 to 10)

    Ensure the response is ONLY valid JSON. No markdown backticks, no other text.
    Category-specific focus: {category_guidelines}
    Title: {title}
    
    Consolidated content to analyze:
    {consolidated_text}
    """
    response = model.invoke([HumanMessage(content=prompt)])
    json_text = response.content.replace('```json', '').replace('```', '').strip()
    return json.loads(json_text)

def run_hierarchical_takeaways_pipeline(transcript, title, category):
    words = transcript.split()
    chunk_size = 15000
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk_words = words[i:i+chunk_size]
        chunks.append(" ".join(chunk_words))
        
    if len(chunks) == 1:
        return generate_final_takeaways_payload(transcript, title, category)
    else:
        with st.spinner("📦 Summarizing video chunks hierarchically..."):
            intermediate_summaries = []
            for idx, chunk in enumerate(chunks):
                st.text(f"Summarizing part {idx+1} of {len(chunks)}...")
                chunk_summary = generate_intermediate_summary(chunk, category)
                intermediate_summaries.append(chunk_summary)
                
            consolidated_text = "\n\n".join(intermediate_summaries)
            return generate_final_takeaways_payload(consolidated_text, title, category)

def create_takeaways_markdown(data, title):
    md = f"# AI Key Takeaways & Executive Summary: {title}\n\n"
    
    md += "## 📝 Executive Summary\n"
    md += "### Short Version (100-200 words)\n"
    md += f"{data['executive_summary']['short']}\n\n"
    md += "### Medium Version (300-500 words)\n"
    md += f"{data['executive_summary']['medium']}\n\n"
    md += "### Detailed Version (600-1000 words)\n"
    md += f"{data['executive_summary']['detailed']}\n\n"
    
    md += "## 🏆 Top 5 Key Takeaways\n"
    for item in data['top_5']:
        md += f"### {item['rank']}. {item['title']}\n"
        md += f"- **Why It Matters**: {item['explanation']}\n"
        md += f"- **How To Apply**: {item['practical_value']}\n\n"
        
    md += "## 📊 Top 10 Key Takeaways (with Importance Scores)\n"
    for item in data['top_10']:
        md += f"### {item['rank']}. {item['title']} (Importance Score: {item['importance_score']}/100)\n"
        md += f"- **Explanation**: {item['explanation']}\n"
        md += f"- **Practical Application**: {item['practical_application']}\n\n"
        
    md += "## 🧠 AI Insight Analytics\n"
    md += f"- **Insights Extracted**: {data['analytics']['num_insights_extracted']}\n"
    md += f"- **Highest Scoring Insight**: {data['analytics']['highest_scoring_insight']}\n"
    md += f"- **Most Discussed Topic**: {data['analytics']['most_discussed_topic']}\n"
    md += f"- **Most Impactful Takeaway**: {data['analytics']['most_impactful_takeaway']}\n"
    md += f"- **Content Complexity Score**: {data['analytics']['content_complexity_score']}/10\n"
    md += f"- **Learning Value Score**: {data['analytics']['learning_value_score']}/10\n"
    
    return md

def create_takeaways_pdf(data, title):
    pdf = BeautifulPDF(f"AI Takeaways: {title}")
    pdf.set_auto_page_break(auto=True, margin=15)
    
    pdf.add_page()
    pdf.set_text_color(75, 0, 130) # Indigo
    pdf.set_font("Helvetica", 'B', 22)
    pdf.ln(60)
    title_clean = title.encode('latin-1', 'replace').decode('latin-1')
    pdf.cell(0, 15, "AI Key Takeaways & Executive Summary", ln=True, align='C')
    pdf.set_font("Helvetica", 'I', 16)
    pdf.cell(0, 10, title_clean, ln=True, align='C')
    pdf.set_text_color(100, 100, 100)
    pdf.set_font("Helvetica", size=14)
    pdf.ln(10)
    pdf.cell(0, 10, "Generated by AI Intelligence System", ln=True, align='C')
    pdf.cell(0, 10, datetime.now().strftime('%B %d, %Y'), ln=True, align='C')
    
    pdf.add_page()
    pdf.set_text_color(75, 0, 130)
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "EXECUTIVE SUMMARY (Detailed)", ln=True)
    pdf.ln(5)
    pdf.set_text_color(30, 30, 30)
    pdf.set_font("Helvetica", size=10)
    summary_text = data['executive_summary']['detailed'].encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 6, summary_text)
    
    pdf.add_page()
    pdf.set_text_color(75, 0, 130)
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "TOP 5 KEY TAKEAWAYS", ln=True)
    pdf.ln(5)
    
    for item in data['top_5']:
        pdf.set_font("Helvetica", 'B', 12)
        pdf.set_text_color(106, 90, 205)
        title_item = f"{item['rank']}. {item['title']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 8, title_item, ln=True)
        
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(30, 30, 30)
        exp_text = f"Why It Matters: {item['explanation']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, exp_text)
        
        val_text = f"How To Apply: {item['practical_value']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, val_text)
        pdf.ln(4)
        
    pdf.add_page()
    pdf.set_text_color(75, 0, 130)
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "TOP 10 KEY TAKEAWAYS (with Scores)", ln=True)
    pdf.ln(5)
    
    for item in data['top_10']:
        pdf.set_font("Helvetica", 'B', 11)
        pdf.set_text_color(106, 90, 205)
        title_item = f"{item['rank']}. {item['title']} (Importance Score: {item['importance_score']}/100)".encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 8, title_item, ln=True)
        
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(30, 30, 30)
        exp_text = f"Explanation: {item['explanation']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, exp_text)
        
        app_text = f"Practical Application: {item['practical_application']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, app_text)
        pdf.ln(4)
        
    return pdf.output(dest='S')

def create_takeaways_docx(data, title):
    doc = Document()
    
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("AI Key Takeaways & Executive Summary")
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_paragraph()
    
    h = doc.add_paragraph()
    run = h.add_run("Executive Summary (Detailed)")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    p = doc.add_paragraph()
    run = p.add_run(data['executive_summary']['detailed'])
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    h = doc.add_paragraph()
    run = h.add_run("Top 5 Key Takeaways")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    for item in data['top_5']:
        p_title = doc.add_paragraph()
        run = p_title.add_run(f"{item['rank']}. {item['title']}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(106, 90, 205)
        
        p_body = doc.add_paragraph()
        run = p_body.add_run(f"Why It Matters: {item['explanation']}\nHow To Apply: {item['practical_value']}")
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        
    doc.add_paragraph()
    
    h = doc.add_paragraph()
    run = h.add_run("Top 10 Key Takeaways (with Scores)")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    for item in data['top_10']:
        p_title = doc.add_paragraph()
        run = p_title.add_run(f"{item['rank']}. {item['title']} (Importance Score: {item['importance_score']}/100)")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(106, 90, 205)
        
        p_body = doc.add_paragraph()
        run = p_body.add_run(f"Explanation: {item['explanation']}\nPractical Application: {item['practical_application']}")
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def takeaways_generator_view():
    st.header("💡 AI Key Takeaways & Executive Summary")
    st.write("Extract deep learning insights, rank key takeaways by strategic importance, and view a structured business-grade executive summary of the content.")
    
    if 'takeaways_transcript' not in st.session_state:
        st.session_state.takeaways_transcript = None
    if 'takeaways_title' not in st.session_state:
        st.session_state.takeaways_title = "Uploaded Media Insights"
    if 'takeaways_category' not in st.session_state:
        st.session_state.takeaways_category = None
    if 'takeaways_payload' not in st.session_state:
        st.session_state.takeaways_payload = {}
        
    source_type = st.radio("Choose Input Source", ["YouTube Video URL", "Upload Audio/Video File"], horizontal=True, key="tk_source")
    
    if source_type == "YouTube Video URL":
        youtube_url = st.text_input("Enter YouTube Video URL", placeholder="https://www.youtube.com/watch?v=...", key="tk_url")
        if st.button("🔍 Get Transcript", key="tk_get_trans"):
            if youtube_url:
                st.info("Fetching video info and transcript...")
                audio_file, info = download_audio(youtube_url)
                if info:
                    st.session_state.takeaways_title = info.get('title', 'YouTube Video')
                
                transcript = get_youtube_transcript(youtube_url)
                if not transcript and audio_file:
                    transcript = transcribe_media_with_gemini(audio_file)
                    
                if audio_file and os.path.exists(audio_file):
                    os.remove(audio_file)
                    
                if transcript:
                    st.session_state.takeaways_transcript = transcript
                    st.success("Transcript generated successfully!")
                    st.rerun()
                else:
                    st.error("Failed to generate transcript.")
    else:
        uploaded_file = st.file_uploader("Upload Audio or Video File", type=["mp3", "m4a", "wav", "mp4", "mov", "avi", "webm"], key="tk_uploader")
        if uploaded_file is not None:
            st.session_state.takeaways_title = uploaded_file.name
            if st.button("🎙️ Transcribe Uploaded File", key="tk_trans_file"):
                if not os.path.exists("temp_audio"):
                    os.makedirs("temp_audio")
                temp_path = os.path.join("temp_audio", uploaded_file.name)
                with open(temp_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                transcript = transcribe_media_with_gemini(temp_path)
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
                if transcript:
                    st.session_state.takeaways_transcript = transcript
                    st.success("File transcribed successfully!")
                    st.rerun()
                else:
                    st.error("Failed to transcribe file.")

    if st.session_state.takeaways_transcript:
        st.divider()
        st.subheader("Transcript Preview")
        with st.expander("Show Full Transcript"):
            st.write(st.session_state.takeaways_transcript)
            
        st.divider()
        st.subheader("Generate Takeaways & Executive Summaries")
        
        if st.button("💡 Extract Executive Takeaways & Summaries", key="tk_generate_all"):
            with st.spinner("🔍 Detecting content category..."):
                st.session_state.takeaways_category = detect_category(st.session_state.takeaways_transcript)
                st.info(f"Category Detected: **{st.session_state.takeaways_category}**")
                
            with st.spinner("🧠 Performing hierarchical extraction and importance scoring..."):
                payload = run_hierarchical_takeaways_pipeline(
                    st.session_state.takeaways_transcript,
                    st.session_state.takeaways_title,
                    st.session_state.takeaways_category
                )
                if payload:
                    st.session_state.takeaways_payload = payload
                    st.success("Takeaways & summaries generated successfully!")
                    st.rerun()
                else:
                    st.error("Failed to generate insights.")

    if st.session_state.takeaways_payload:
        data = st.session_state.takeaways_payload
        
        st.divider()
        st.subheader("📊 AI Insight Analytics Dashboard")
        
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric(label="Insights Extracted", value=f"{data['analytics']['num_insights_extracted']}")
        with c2:
            st.metric(label="Content Complexity Score", value=f"{data['analytics']['content_complexity_score']}/10")
        with c3:
            st.metric(label="Learning Value Score", value=f"{data['analytics']['learning_value_score']}/10")
            
        st.markdown(f"**Most Discussed Topic:** `{data['analytics']['most_discussed_topic']}`")
        st.markdown(f"**Highest Scoring Insight:** *\"{data['analytics']['highest_scoring_insight']}\"*")
        st.markdown(f"**Most Impactful Takeaway:** *\"{data['analytics']['most_impactful_takeaway']}\"*")
        
        st.divider()
        st.subheader("📋 Insights & Summary Sections")
        
        t1, t2, t3 = st.tabs(["🏆 Top 5 Takeaways", "📊 Top 10 Takeaways", "📝 Executive Summary"])
        
        with t1:
            for item in data['top_5']:
                st.markdown(
                    f"""
                    <div style="background-color:#FFFFFF; padding:20px; border-radius:15px; margin-bottom:15px; border-left: 5px solid #9370DB; box-shadow: 0 4px 6px rgba(0,0,0,0.05);">
                        <h4 style="color:#4B0082; margin-top:0;">{item['rank']}. {item['title']}</h4>
                        <p><b>Why It Matters:</b> {item['explanation']}</p>
                        <p style="color:#6A5ACD;"><b>How to Apply:</b> {item['practical_value']}</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                
        with t2:
            st.write("### Filter & Sort Takeaways")
            col_f1, col_f2, col_f3 = st.columns([1.5, 1, 1])
            with col_f1:
                search_query = st.text_input("Search Takeaways", "", placeholder="Search keywords...")
            with col_f2:
                min_score = st.slider("Min Importance Score", 0, 100, 0)
            with col_f3:
                sort_order = st.selectbox("Sort Order", ["Highest Score", "Relevance/Rank", "Lowest Score"])
                
            filtered_takeaways = []
            for item in data['top_10']:
                matches_search = (
                    search_query.lower() in item['title'].lower() or
                    search_query.lower() in item['explanation'].lower() or
                    search_query.lower() in item['practical_application'].lower()
                )
                matches_score = item['importance_score'] >= min_score
                if matches_search and matches_score:
                    filtered_takeaways.append(item)
                    
            if sort_order == "Highest Score":
                filtered_takeaways.sort(key=lambda x: x['importance_score'], reverse=True)
            elif sort_order == "Lowest Score":
                filtered_takeaways.sort(key=lambda x: x['importance_score'])
            else:
                filtered_takeaways.sort(key=lambda x: x['rank'])
                
            for item in filtered_takeaways:
                st.markdown(
                    f"""
                    <div style="background-color:#FFFFFF; padding:20px; border-radius:15px; margin-bottom:15px; border-left: 5px solid #6A5ACD; box-shadow: 0 4px 6px rgba(0,0,0,0.05);">
                        <div style="display:flex; justify-content:space-between; align-items:center;">
                            <h4 style="color:#4B0082; margin:0;">{item['rank']}. {item['title']}</h4>
                            <span style="background-color:#E6E6FA; color:#4B0082; padding:4px 8px; border-radius:10px; font-weight:bold; font-size:12px;">Score: {item['importance_score']}/100</span>
                        </div>
                        <p style="margin-top:10px;"><b>Explanation:</b> {item['explanation']}</p>
                        <p style="color:#6A5ACD;"><b>Practical Application:</b> {item['practical_application']}</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            if not filtered_takeaways:
                st.info("No takeaways match your filters.")
                
        with t3:
            len_choice = st.radio("Select Summary Length", ["Short (100-200 words)", "Medium (300-500 words)", "Detailed (600-1000 words)"], horizontal=True)
            
            st.divider()
            if len_choice.startswith("Short"):
                st.write(data['executive_summary']['short'])
            elif len_choice.startswith("Medium"):
                st.write(data['executive_summary']['medium'])
            else:
                st.write(data['executive_summary']['detailed'])
                
        st.divider()
        st.subheader("📥 Export & Download Takeaways")
        
        c_ex1, c_ex2, c_ex3, c_ex4 = st.columns(4)
        markdown_content = create_takeaways_markdown(data, st.session_state.takeaways_title)
        
        with c_ex1:
            st.download_button(
                label="📄 Markdown (.md)",
                data=markdown_content,
                file_name=f"{st.session_state.takeaways_title.replace(' ', '_')}_Takeaways.md",
                mime="text/markdown",
                key="dl_tk_md"
            )
        with c_ex2:
            pdf_bytes = create_takeaways_pdf(data, st.session_state.takeaways_title)
            st.download_button(
                label="📕 PDF Report",
                data=pdf_bytes,
                file_name=f"{st.session_state.takeaways_title.replace(' ', '_')}_Takeaways.pdf",
                mime="application/pdf",
                key="dl_tk_pdf"
            )
        with c_ex3:
            docx_bytes = create_takeaways_docx(data, st.session_state.takeaways_title)
            st.download_button(
                label="📘 Word (.docx)",
                data=docx_bytes,
                file_name=f"{st.session_state.takeaways_title.replace(' ', '_')}_Takeaways.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_tk_docx"
            )
        with c_ex4:
            st.download_button(
                label="📝 Plain Text (.txt)",
                data=markdown_content,
                file_name=f"{st.session_state.takeaways_title.replace(' ', '_')}_Takeaways.txt",
                mime="text/plain",
                key="dl_tk_txt"
            )

def get_youtube_transcript_with_timestamps(url):
    video_id = None
    patterns = [
        r"(?:v=|\/v\/|\/embed\/|\/shorts\/|youtu\.be\/)([a-zA-Z0-9_-]{11})"
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            video_id = match.group(1)
            break
            
    if not video_id:
        return None, None
        
    try:
        print("VIDEO ID:", video_id)
        print("TRANSCRIPT API START")
        
        try:
            api_instance = YouTubeTranscriptApi()
            transcript_list_obj = api_instance.list(video_id)
            transcript_obj = transcript_list_obj.find_transcript(['en', 'hi', 'te', 'ta', 'kn'])
            raw_transcript = transcript_obj.fetch()
            
            transcript_list = []
            for item in raw_transcript:
                transcript_list.append({
                    'text': getattr(item, 'text', item.get('text', '')),
                    'start': getattr(item, 'start', item.get('start', 0.0))
                })
        except AttributeError:
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en', 'hi', 'te', 'ta', 'kn'])
            
        print("TRANSCRIPT RETRIEVED")
        timed_segments = []
        for item in transcript_list:
            start = int(item['start'])
            h = start // 3600
            m = (start % 3600) // 60
            s = start % 60
            ts_str = f"[{h:02d}:{m:02d}:{s:02d}]" if h > 0 else f"[{m:02d}:{s:02d}]"
            timed_segments.append(f"{ts_str} {item['text']}")
        return "\n".join(timed_segments), transcript_list
    except Exception as e:
        print("TRANSCRIPT FAILED:", str(e))
        print("WHISPER FALLBACK STARTED")
        return None, None

def get_youtube_original_chapters(info):
    if not info:
        return []
        
    chapters = info.get('chapters')
    if chapters:
        formatted_chapters = []
        for c in chapters:
            start = int(c.get('start_time', 0))
            h = start // 3600
            m = (start % 3600) // 60
            s = start % 60
            ts_str = f"{h:02d}:{m:02d}:{s:02d}" if h > 0 else f"{m:02d}:{s:02d}"
            formatted_chapters.append({
                "start_time": ts_str,
                "seconds": start,
                "title": c.get('title', 'Chapter'),
                "summary": "Original video chapter (No summary available).",
                "importance_score": 80,
                "duration_seconds": int(c.get('end_time', start) - start),
                "concepts": []
            })
        return formatted_chapters

    desc = info.get('description', '')
    parsed = []
    pattern = r"((?:\d{1,2}:)?\d{1,2}:\d{2})\s+(.+)"
    for line in desc.split('\n'):
        match = re.search(pattern, line)
        if match:
            ts_str = match.group(1)
            title = match.group(2).strip(" -|")
            parts = list(map(int, ts_str.split(':')))
            if len(parts) == 2:
                seconds = parts[0] * 60 + parts[1]
            elif len(parts) == 3:
                seconds = parts[0] * 3600 + parts[1] * 60 + parts[2]
            else:
                seconds = 0
            parsed.append({
                "start_time": ts_str,
                "seconds": seconds,
                "title": title,
                "summary": "Original parsed chapter (No summary available).",
                "importance_score": 80,
                "duration_seconds": 0,
                "concepts": []
            })
            
    for idx in range(len(parsed) - 1):
        parsed[idx]["duration_seconds"] = parsed[idx+1]["seconds"] - parsed[idx]["seconds"]
        
    return parsed

def generate_ai_chapters(timestamped_transcript, length_choice):
    model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
    if length_choice == "Short":
        summary_guidelines = "Summaries must be concise (1–2 sentences)."
    elif length_choice == "Medium":
        summary_guidelines = "Summaries must be clear and informative (3–5 sentences)."
    else:
        summary_guidelines = "Summaries must be detailed and comprehensive (5–10 sentences)."
        
    prompt = f"""
    You are an expert video editor and NLP analyst. Analyze the following timestamped transcript of a video and perform topic segmentation.
    Identify natural chapter boundaries based on semantic shifts, topic changes, and content transitions. Do not generate timestamps at fixed intervals.
    
    For each chapter detected, you must return:
    - start_time: string in format "MM:SS" or "HH:MM:SS" matching the closest timestamp in the transcript.
    - seconds: integer representing the start time in seconds.
    - title: string (descriptive, meaningful, 3-8 words. Avoid generic labels like Chapter 1, Section 2).
    - summary: string (summarize the chapter contents. {summary_guidelines} Mention tools, examples, or decisions discussed).
    - importance_score: integer (1 to 100, based on impact and relevance).
    - duration_seconds: integer (estimated duration of this chapter in seconds).
    - concepts: list of strings (key concepts, definitions, or frameworks discussed in this chapter).
    
    Provide the output in valid JSON format as a list of dicts.
    Return ONLY the raw JSON. No markdown backticks, no other text.
    
    Transcript with Timestamps:
    {timestamped_transcript[:25000]}
    """
    response = model.invoke([HumanMessage(content=prompt)])
    json_text = response.content.replace('```json', '').replace('```', '').strip()
    return json.loads(json_text)

def merge_chunk_chapters(chunk_chapters_list, title):
    flat_chapters = []
    for chunk in chunk_chapters_list:
        for ch in chunk:
            flat_chapters.append(ch)
            
    model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
    prompt = f"""
    You are an expert editor. You have been given a list of chapters generated from different chunks of a video transcript.
    Some adjacent chapters might cover duplicate topics or represent redundant divisions.
    Merge adjacent chapters that are semantically similar into a single cohesive chapter, ensuring the timeline flows logically.
    
    For the merged list, return a valid JSON array of objects with keys:
    - start_time: string (e.g. "MM:SS" or "HH:MM:SS")
    - seconds: integer
    - title: string
    - summary: string
    - importance_score: integer
    - duration_seconds: integer
    - concepts: list of strings
    
    Return ONLY the raw JSON. No markdown backticks, no other text.
    
    Video Title: {title}
    Chapters to consolidate:
    {json.dumps(flat_chapters, indent=2)}
    """
    response = model.invoke([HumanMessage(content=prompt)])
    json_text = response.content.replace('```json', '').replace('```', '').strip()
    return json.loads(json_text)

def generate_hybrid_chapters(original, ai):
    if not original:
        return ai
    if not ai:
        return original
        
    model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
    prompt = f"""
    You are an expert editor. Merge the following original video chapters and AI-generated chapters.
    Use the precise timestamps of the original chapters where appropriate, but enrich them with the descriptive titles, detailed summaries, importance scores, and key concepts from the AI chapters.
    
    Provide the output in valid JSON format as a list of dicts with keys:
    - start_time: string (e.g. "MM:SS" or "HH:MM:SS")
    - seconds: integer
    - title: string
    - summary: string
    - importance_score: integer
    - duration_seconds: integer
    - concepts: list of strings
    
    Return ONLY the raw JSON. No markdown backticks, no other text.
    
    Original Chapters:
    {json.dumps(original, indent=2)}
    
    AI Chapters:
    {json.dumps(ai, indent=2)}
    """
    response = model.invoke([HumanMessage(content=prompt)])
    json_text = response.content.replace('```json', '').replace('```', '').strip()
    return json.loads(json_text)

def run_chapters_pipeline(timestamped_transcript, title, length_choice, url, info):
    words = timestamped_transcript.split()
    chunk_size = 15000
    chunks = []
    
    for i in range(0, len(words), chunk_size):
        chunk_words = words[i:i+chunk_size]
        chunks.append(" ".join(chunk_words))
        
    ai_chapters = []
    
    if len(chunks) == 1:
        ai_chapters = generate_ai_chapters(timestamped_transcript, length_choice)
    else:
        with st.spinner("📦 Generating chapter segments hierarchically..."):
            chunk_results = []
            for j, chunk in enumerate(chunks):
                st.text(f"Analyzing chunk {j+1} of {len(chunks)}...")
                chunk_res = generate_ai_chapters(chunk, length_choice)
                chunk_results.append(chunk_res)
                
            ai_chapters = merge_chunk_chapters(chunk_results, title)
            
    original_chapters = get_youtube_original_chapters(info)
    
    hybrid_chapters = []
    if original_chapters and ai_chapters:
        with st.spinner("⚖️ Merging original & AI chapters to generate Hybrid Timeline..."):
            hybrid_chapters = generate_hybrid_chapters(original_chapters, ai_chapters)
            
    return {
        "original": original_chapters,
        "ai": ai_chapters,
        "hybrid": hybrid_chapters if hybrid_chapters else ai_chapters
    }

def create_chapters_markdown(chapters, title, mode="Standard"):
    md = f"# AI Video Chapters & Timeline: {title}\n\n"
    md += f"Timeline Mode: **{mode}**\n\n"
    
    for idx, c in enumerate(chapters):
        md += f"### {c['start_time']} - {c['title']} (Importance: {c['importance_score']}/100)\n"
        if mode in ["Standard", "Detailed", "Learning"]:
            md += f"- **Summary**: {c['summary']}\n"
            
        if mode == "Detailed":
            if c.get('concepts'):
                md += f"- **Key Concepts**: {', '.join(c['concepts'])}\n"
            md += f"- **Duration**: {c['duration_seconds']}s\n"
            
        if mode == "Learning":
            if c.get('concepts'):
                md += f"- **Key Concepts**: {', '.join(c['concepts'])}\n"
            md += f"- **Key Takeaway**: Understand core concept of {c['title']}.\n"
            md += f"- **Revision Point**: Remember that: {c['summary'][:100]}...\n"
        md += "\n"
        
    return md

def create_chapters_pdf(chapters, title, mode="Standard"):
    pdf = BeautifulPDF(f"AI Chapters: {title}")
    pdf.set_auto_page_break(auto=True, margin=15)
    
    pdf.add_page()
    pdf.set_text_color(75, 0, 130) # Indigo
    pdf.set_font("Helvetica", 'B', 22)
    pdf.ln(60)
    title_clean = title.encode('latin-1', 'replace').decode('latin-1')
    pdf.cell(0, 15, "AI Timestamp Chapters & Timeline", ln=True, align='C')
    pdf.set_font("Helvetica", 'I', 16)
    pdf.cell(0, 10, title_clean, ln=True, align='C')
    pdf.set_text_color(100, 100, 100)
    pdf.set_font("Helvetica", size=14)
    pdf.ln(10)
    pdf.cell(0, 10, "Generated by AI Video Breakdown System", ln=True, align='C')
    pdf.cell(0, 10, datetime.now().strftime('%B %d, %Y'), ln=True, align='C')
    
    pdf.add_page()
    pdf.set_text_color(75, 0, 130)
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "VIDEO CHAPTERS TIMELINE", ln=True)
    pdf.ln(5)
    
    for c in chapters:
        pdf.set_font("Helvetica", 'B', 12)
        pdf.set_text_color(106, 90, 205)
        title_item = f"[{c['start_time']}] {c['title']} (Importance: {c['importance_score']}/100)".encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 8, title_item, ln=True)
        
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(30, 30, 30)
        if mode in ["Standard", "Detailed", "Learning"]:
            summary_clean = f"Summary: {c['summary']}".encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, summary_clean)
            
        if mode in ["Detailed", "Learning"] and c.get('concepts'):
            concepts_clean = f"Concepts: {', '.join(c['concepts'])}".encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, concepts_clean)
        pdf.ln(4)
        
    return pdf.output(dest='S')

def create_chapters_docx(chapters, title, mode="Standard"):
    doc = Document()
    
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("AI Timestamp Chapters & Timeline")
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_paragraph()
    
    for c in chapters:
        p_title = doc.add_paragraph()
        run = p_title.add_run(f"[{c['start_time']}] {c['title']} (Importance Score: {c['importance_score']}/100)")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(106, 90, 205)
        
        if mode in ["Standard", "Detailed", "Learning"]:
            p_body = doc.add_paragraph()
            run = p_body.add_run(f"Summary: {c['summary']}")
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            
        if mode in ["Detailed", "Learning"] and c.get('concepts'):
            p_concepts = doc.add_paragraph()
            run = p_concepts.add_run(f"Concepts: {', '.join(c['concepts'])}")
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.italic = True
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_chapters_csv(chapters):
    output = io.StringIO()
    output.write("Timestamp,Chapter Title,Summary,Importance Score,Topic Metadata\n")
    for c in chapters:
        title = c['title'].replace('"', '""')
        summary = c['summary'].replace('"', '""')
        concepts = ", ".join(c.get('concepts', [])).replace('"', '""')
        output.write(f'"{c["start_time"]}","{title}","{summary}","{c["importance_score"]}","{concepts}"\n')
    return output.getvalue()

def create_chapters_json(chapters):
    return json.dumps(chapters, indent=2)

def timestamp_summary_view():
    st.header("⏰ AI Timestamp Summary & Chapter Breakdown")
    st.write("Generate interactive timelines, chapter segments, and summary breakdowns for any video or uploaded audio/video file.")
    
    if 'ts_transcript' not in st.session_state:
        st.session_state.ts_transcript = None
    if 'ts_list' not in st.session_state:
        st.session_state.ts_list = None
    if 'ts_title' not in st.session_state:
        st.session_state.ts_title = "Uploaded Media Timeline"
    if 'ts_chapters' not in st.session_state:
        st.session_state.ts_chapters = {}
    if 'ts_url' not in st.session_state:
        st.session_state.ts_url = None
    if 'ts_info' not in st.session_state:
        st.session_state.ts_info = None
    if 'video_jump_time' not in st.session_state:
        st.session_state.video_jump_time = 0
        
    source_type = st.radio("Choose Input Source", ["YouTube Video URL", "Upload Audio/Video File"], horizontal=True, key="ts_source")
    
    if source_type == "YouTube Video URL":
        youtube_url = st.text_input("Enter YouTube Video URL", placeholder="https://www.youtube.com/watch?v=...", key="ts_input_url")
        if st.button("🔍 Get Transcript & Timestamps", key="ts_get_trans"):
            if youtube_url:
                st.info("Fetching video info, transcript, and timestamps...")
                audio_file, info = download_audio(youtube_url)
                if info:
                    st.session_state.ts_title = info.get('title', 'YouTube Video')
                    st.session_state.ts_info = info
                    st.session_state.ts_url = youtube_url
                    
                transcript_text, transcript_list = get_youtube_transcript_with_timestamps(youtube_url)
                if not transcript_text and audio_file:
                    transcript_text = transcribe_media_with_gemini(audio_file)
                    
                if audio_file and os.path.exists(audio_file):
                    os.remove(audio_file)
                    
                if transcript_text:
                    st.session_state.ts_transcript = transcript_text
                    st.session_state.ts_list = transcript_list
                    st.success("Timestamps and Transcript generated successfully!")
                    st.rerun()
                else:
                    st.error("Failed to generate transcript.")
    else:
        uploaded_file = st.file_uploader("Upload Audio or Video File", type=["mp3", "m4a", "wav", "mp4", "mov", "avi", "webm"], key="ts_uploader")
        if uploaded_file is not None:
            st.session_state.ts_title = uploaded_file.name
            st.session_state.ts_url = None
            if st.button("🎙️ Transcribe Uploaded File", key="ts_trans_file"):
                if not os.path.exists("temp_audio"):
                    os.makedirs("temp_audio")
                temp_path = os.path.join("temp_audio", uploaded_file.name)
                with open(temp_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                transcript_text = transcribe_media_with_gemini(temp_path)
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
                if transcript_text:
                    st.session_state.ts_transcript = transcript_text
                    st.success("File transcribed successfully!")
                    st.rerun()
                else:
                    st.error("Failed to transcribe file.")

    if st.session_state.ts_transcript:
        st.divider()
        st.subheader("Generate Timelines & Video Chapters")
        
        c_len, c_type = st.columns(2)
        with c_len:
            summary_len = st.selectbox("Select Chapter Summary Length", ["Short (1-2 sentences)", "Medium (3-5 sentences)", "Detailed (5-10 sentences)"])
        with c_type:
            pass
            
        if st.button("🧠 Segment & Breakdown Video", key="ts_segment"):
            len_choice = "Short" if "Short" in summary_len else "Medium" if "Medium" in summary_len else "Detailed"
            payloads = run_chapters_pipeline(
                st.session_state.ts_transcript,
                st.session_state.ts_title,
                len_choice,
                st.session_state.ts_url,
                st.session_state.ts_info
            )
            if payloads:
                st.session_state.ts_chapters = payloads
                st.success("Video breakdown complete!")
                st.rerun()
            else:
                st.error("Failed to generate segments.")

    if st.session_state.ts_chapters:
        payload = st.session_state.ts_chapters
        
        st.divider()
        st.subheader("Timeline Integration Selector")
        avail_options = ["AI Generated Chapters"]
        if payload.get("original"):
            avail_options.append("Original Video Chapters")
            avail_options.append("Hybrid Chapters (AI Enriched)")
            
        timeline_choice = st.selectbox("Choose Timeline Integration Source:", avail_options, index=0)
        
        chapters = []
        if timeline_choice == "Original Video Chapters":
            chapters = payload.get("original", [])
        elif timeline_choice == "Hybrid Chapters (AI Enriched)":
            chapters = payload.get("hybrid", [])
        else:
            chapters = payload.get("ai", [])
            
        st.divider()
        
        if chapters:
            st.subheader("📊 Chapter Analytics")
            total_chapters = len(chapters)
            total_sec = sum(c.get('duration_seconds', 0) for c in chapters)
            avg_sec = total_sec // total_chapters if total_chapters > 0 else 0
            
            most_imp = max(chapters, key=lambda x: x.get('importance_score', 0))
            
            col_an1, col_an2, col_an3 = st.columns(3)
            with col_an1:
                st.metric(label="Total Chapters", value=total_chapters)
            with col_an2:
                st.metric(label="Avg Chapter Length", value=f"{avg_sec // 60}m {avg_sec % 60}s")
            with col_an3:
                st.metric(label="Most Important Chapter", value=f"[{most_imp['start_time']}]")
            st.write(f"**Highest Scoring Chapter Title:** *\"{most_imp['title']}\"* (Importance: {most_imp['importance_score']}/100)")
            
            st.divider()
            
            col_sch1, col_sch2 = st.columns([2, 1])
            with col_sch1:
                ch_query = st.text_input("Search Timeline", "", placeholder="Search titles, concepts, keywords...", key="ts_ch_search")
            with col_sch2:
                min_ch_score = st.slider("Min Chapter Importance", 0, 100, 0, key="ts_ch_score")
                
            filtered_chapters = []
            for c in chapters:
                matches_search = (
                    ch_query.lower() in c['title'].lower() or
                    ch_query.lower() in c['summary'].lower() or
                    any(ch_query.lower() in con.lower() for con in c.get('concepts', []))
                )
                matches_score = c['importance_score'] >= min_ch_score
                if matches_search and matches_score:
                    filtered_chapters.append(c)
                    
            st.divider()
            
            timeline_mode = st.radio("Select Timeline UI Mode", ["Quick (Titles Only)", "Standard (Summaries)", "Detailed (with Concepts)", "Learning (with Takeaways)"], horizontal=True)
            
            col_v1, col_v2 = st.columns([1.2, 1])
            
            with col_v1:
                st.subheader("🎥 Video Player")
                if st.session_state.ts_url:
                    st.video(st.session_state.ts_url, start_time=st.session_state.video_jump_time)
                else:
                    st.warning("No video URL available for playback (Local upload).")
                    
                st.info("💡 Tip: Click on any timestamp in the interactive timeline on the right to jump the video directly to that chapter!")
                
            with col_v2:
                st.subheader("🧭 Interactive Timeline")
                
                for idx, c in enumerate(filtered_chapters):
                    c_title = c['title']
                    c_sec = c['seconds']
                    c_ts = c['start_time']
                    
                    st.markdown(
                        f"""
                        <div style="background-color:#FFFFFF; padding:15px; border-radius:15px; margin-bottom:10px; border-left: 5px solid #6A5ACD; box-shadow: 0 4px 6px rgba(0,0,0,0.05);">
                            <div style="display:flex; justify-content:space-between; align-items:center;">
                                <span style="font-weight:bold; color:#4B0082; font-size:16px;">⏱️ {c_ts}</span>
                                <span style="background-color:#E6E6FA; color:#4B0082; padding:2px 6px; border-radius:8px; font-weight:bold; font-size:11px;">Score: {c['importance_score']}/100</span>
                            </div>
                            <h4 style="margin:5px 0 10px 0; color:#4B0082;">{c_title}</h4>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                    
                    col_btn, col_blank = st.columns([1.2, 2])
                    with col_btn:
                        if st.button(f"Jump to {c_ts}", key=f"jmp_{c_ts}_{idx}"):
                            st.session_state.video_jump_time = c_sec
                            st.rerun()
                            
                    if timeline_mode != "Quick (Titles Only)":
                        st.write(c['summary'])
                        
                    if timeline_mode in ["Detailed (with Concepts)", "Learning (with Takeaways)"] and c.get('concepts'):
                        st.write(f"**Key Concepts:** {', '.join(c['concepts'])}")
                        
                    if timeline_mode == "Learning (with Takeaways)":
                        st.markdown(f"- 💡 **Key Takeaway**: Understand core concept of {c_title}.")
                        st.markdown(f"- 📝 **Revision Point**: Remember that: {c['summary'][:120]}...")
                    st.divider()
                    
            st.divider()
            st.subheader("📥 Export & Download Timeline")
            
            export_mode = "Quick" if "Quick" in timeline_mode else "Standard" if "Standard" in timeline_mode else "Detailed" if "Detailed" in timeline_mode else "Learning"
            
            md_timeline = create_chapters_markdown(chapters, st.session_state.ts_title, export_mode)
            pdf_timeline = create_chapters_pdf(chapters, st.session_state.ts_title, export_mode)
            docx_timeline = create_chapters_docx(chapters, st.session_state.ts_title, export_mode)
            csv_timeline = create_chapters_csv(chapters)
            json_timeline = create_chapters_json(chapters)
            
            col_ex1, col_ex2, col_ex3, col_ex4, col_ex5 = st.columns(5)
            with col_ex1:
                st.download_button(
                    label="📄 Markdown (.md)",
                    data=md_timeline,
                    file_name=f"{st.session_state.ts_title.replace(' ', '_')}_Chapters.md",
                    mime="text/markdown",
                    key="dl_ch_md"
                )
            with col_ex2:
                st.download_button(
                    label="📕 PDF Report",
                    data=pdf_timeline,
                    file_name=f"{st.session_state.ts_title.replace(' ', '_')}_Chapters.pdf",
                    mime="application/pdf",
                    key="dl_ch_pdf"
                )
            with col_ex3:
                st.download_button(
                    label="📘 Word (.docx)",
                    data=docx_timeline,
                    file_name=f"{st.session_state.ts_title.replace(' ', '_')}_Chapters.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_ch_docx"
                )
            with col_ex4:
                st.download_button(
                    label="📊 CSV Table",
                    data=csv_timeline,
                    file_name=f"{st.session_state.ts_title.replace(' ', '_')}_Chapters.csv",
                    mime="text/csv",
                    key="dl_ch_csv"
                )
            with col_ex5:
                st.download_button(
                    label="⚙️ JSON Data",
                    data=json_timeline,
                    file_name=f"{st.session_state.ts_title.replace(' ', '_')}_Chapters.json",
                    mime="application/json",
                    key="dl_ch_json"
                )

def create_chat_pdf(history, title):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", 'B', 16)
        pdf.cell(0, 10, f"Chat History - {title}", ln=True, align='C')
        pdf.ln(10)
        
        for msg in history:
            role = "USER" if msg["role"] == "user" else "AI ASSISTANT"
            pdf.set_font("Helvetica", 'B', 11)
            pdf.cell(0, 8, f"{role}:", ln=True)
            pdf.set_font("Helvetica", size=10)
            
            clean_content = msg["content"].encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, clean_content)
            pdf.ln(4)
            
        return pdf.output()
    except Exception as e:
        st.error(f"Error creating PDF: {e}")
        return None

def create_chat_docx(history, title):
    try:
        doc = Document()
        doc.add_heading(f"Chat History - {title}", 0)
        
        for msg in history:
            role = "User" if msg["role"] == "user" else "AI Assistant"
            p = doc.add_paragraph()
            run = p.add_run(f"{role}: ")
            run.bold = True
            p.add_run(msg["content"])
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"Error creating DOCX: {e}")
        return None

def create_chat_markdown(history, title):
    lines = [f"# Chat History - {title}\n"]
    for msg in history:
        role = "**User**" if msg["role"] == "user" else "**AI Assistant**"
        lines.append(f"{role}:\n{msg['content']}\n\n---\n")
    return "\n".join(lines)

def create_chat_json(history):
    return json.dumps(history, indent=2)

def rag_chatbot_view():
    st.header("💬 AI Question Answering Chatbot (RAG)")
    st.write("Interact conversationally with indexed videos. All answers are grounded in the transcript with timestamped citations.")
    
    if 'rag_store' not in st.session_state:
        st.session_state.rag_store = VideoRAGStore(path="chroma_db_storage")
    if 'rag_history' not in st.session_state:
        st.session_state.rag_history = []
    if 'rag_mode' not in st.session_state:
        st.session_state.rag_mode = "Student Mode"
        
    store = st.session_state.rag_store
    indexed_videos = store.get_indexed_videos()
    
    if not indexed_videos:
        st.warning("⚠️ No videos indexed in the database yet. Please process a video in the Home view or get a transcript in the Notes view first.")
        return
        
    st.sidebar.divider()
    st.sidebar.subheader("💬 Chatbot Context Settings")
    
    video_options = {v["video_id"]: v["video_title"] for v in indexed_videos}
    
    st.sidebar.write("**Select Indexed Videos to Search:**")
    selected_vids = []
    for vid_id, vid_title in video_options.items():
        if st.sidebar.checkbox(vid_title, value=True, key=f"sel_chat_{vid_id}"):
            selected_vids.append(vid_id)
            
    st.sidebar.subheader("🧠 Learning Assistant Mode")
    st.session_state.rag_mode = st.sidebar.selectbox(
        "Format responses as:",
        ["Student Mode", "Explain Simply", "Explain Like I'm 10", "Explain Technically", "Interview Mode", "Exam Mode", "Developer Mode", "Research Mode"],
        index=0
    )
    
    st.sidebar.divider()
    st.sidebar.subheader("🗑️ Manage RAG Database")
    del_vid = st.sidebar.selectbox("Select Video to Delete:", ["-- Select Video --"] + list(video_options.values()))
    if st.sidebar.button("🗑️ Delete Selected Video Index"):
        if del_vid != "-- Select Video --":
            vid_to_del = next((vid_id for vid_id, title in video_options.items() if title == del_vid), None)
            if vid_to_del:
                store.delete_video(vid_to_del)
                st.sidebar.success(f"Deleted index for: {del_vid}")
                st.rerun()

    st.subheader("🔍 Fast Search within Transcripts")
    search_col1, search_col2 = st.columns([4, 1])
    with search_col1:
        keyword = st.text_input("Look up keywords, tools, or concepts in transcript (e.g. 'TensorFlow'):", placeholder="Enter query...", key="kw_search_input")
    with search_col2:
        st.write("")
        st.write("")
        run_search = st.button("🔍 Search")
        
    if keyword and (run_search or st.session_state.get("kw_search_input")):
        st.write("**Search Results:**")
        matches = []
        for vid_id in selected_vids:
            matches.extend(store.keyword_search(keyword, vid_id))
            
        if matches:
            for match in matches[:5]:
                meta = match["metadata"]
                title_label = meta.get("video_title", "Video")
                with st.expander(f"⏰ {meta.get('start_time', '00:00')} - {meta.get('end_time', '00:00')} in Chapter: {meta.get('chapter_name', 'N/A')} ({title_label})"):
                    st.write(match["text"])
        else:
            st.info("No matching snippets found.")
            
    st.divider()
    st.subheader("💬 Chat Window")
    
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    for msg in st.session_state.rag_history:
        role = "user-message" if msg['role'] == 'user' else "ai-message"
        avatar = "👤" if msg['role'] == 'user' else "🤖"
        st.markdown(f'<div class="chat-message {role}"><b>{avatar} {role.replace("-message", "").upper()}:</b><br>{msg["content"]}</div>', unsafe_allow_html=True)
        
        if msg['role'] == 'assistant' and 'sources' in msg and msg['sources']:
            with st.expander(f"🎯 Citations ({len(msg['sources'])} sources matched | Confidence: {msg.get('confidence_score', 85)}%)"):
                for idx, src in enumerate(msg['sources']):
                    meta = src["metadata"]
                    st.markdown(f"**Source {idx+1}**: Chapter: `{meta.get('chapter_name', 'N/A')}` | Timestamp: `{meta.get('start_time', '00:00')} - {meta.get('end_time', '00:00')}` | Match Similarity: `{src.get('confidence_score', 90)}%`")
                    st.caption(f'"{src["text"]}"')
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.write("")
    st.write("**Suggested Questions:**")
    sugs = ["What is the main topic of this video?", "What are the key advantages mentioned?", "What tools/libraries were used?", "Can you explain the main workflows?"]
    cols = st.columns(len(sugs))
    clicked_sug = None
    for idx, sug in enumerate(sugs):
        with cols[idx]:
            if st.button(sug, key=f"sug_btn_{idx}"):
                clicked_sug = sug
                
    user_query = st.chat_input("Ask a question about the video transcripts...")
    query_to_run = user_query or clicked_sug
    
    if query_to_run:
        st.session_state.rag_history.append({"role": "user", "content": query_to_run})
        
        if not selected_vids:
            st.session_state.rag_history.append({
                "role": "assistant",
                "content": "⚠️ Please select at least one indexed video from the sidebar checklist to search context from.",
                "sources": []
            })
            st.rerun()
            
        with st.spinner("🧠 Retrieval-Augmented Thinking..."):
            q_emb = generate_embedding(query_to_run)
            contexts = store.query(q_emb, video_ids=selected_vids, top_k=5)
            answer = generate_rag_answer(query_to_run, st.session_state.rag_history[:-1], contexts, st.session_state.rag_mode)
            avg_conf = int(sum(c["confidence_score"] for c in contexts) / len(contexts)) if contexts else 80
            
            st.session_state.rag_history.append({
                "role": "assistant",
                "content": answer,
                "sources": contexts,
                "confidence_score": avg_conf
            })
            st.rerun()
            
    if st.session_state.rag_history:
        st.divider()
        st.subheader("📥 Export Chat History")
        col_dl1, col_dl2, col_dl3, col_dl4 = st.columns(4)
        chat_title = "RAG_Chat"
        
        with col_dl1:
            chat_pdf = create_chat_pdf(st.session_state.rag_history, chat_title)
            if chat_pdf:
                st.download_button("📕 Download PDF", chat_pdf, file_name=f"{chat_title}.pdf", mime="application/pdf", key="dl_chat_pdf")
        with col_dl2:
            chat_docx = create_chat_docx(st.session_state.rag_history, chat_title)
            if chat_docx:
                st.download_button("📘 Download Word (.docx)", chat_docx, file_name=f"{chat_title}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_chat_docx")
        with col_dl3:
            chat_md = create_chat_markdown(st.session_state.rag_history, chat_title)
            st.download_button("📄 Download Markdown (.md)", chat_md, file_name=f"{chat_title}.md", mime="text/markdown", key="dl_chat_md")
        with col_dl4:
            chat_json = create_chat_json(st.session_state.rag_history)
            st.download_button("⚙️ Download JSON Data", chat_json, file_name=f"{chat_title}.json", mime="application/json", key="dl_chat_json")

def mind_map_view():
    st.header("🌳 AI Mind Map & Knowledge Graph Generator")
    st.write("Visualize video transcripts as Concept Trees, Topic Relationships, and Knowledge Graphs.")
    
    if 'rag_store' not in st.session_state:
        st.session_state.rag_store = VideoRAGStore(path="chroma_db_storage")
    if 'active_graph' not in st.session_state:
        st.session_state.active_graph = None
        
    store = st.session_state.rag_store
    indexed_videos = store.get_indexed_videos()
    
    if not indexed_videos:
        st.warning("⚠️ No videos indexed in the database yet. Please process a video in the Home view or get a transcript in the Notes view first.")
        return
        
    col_opt1, col_opt2, col_opt3 = st.columns(3)
    video_options = {v["video_id"]: v["video_title"] for v in indexed_videos}
    
    with col_opt1:
        selected_vid_id = st.selectbox("Select Video to Map:", list(video_options.keys()), format_func=lambda x: video_options[x])
    with col_opt2:
        vis_layout = st.selectbox("Default Visualization Layout:", ["network", "hierarchical"], format_func=lambda x: "Force Directed" if x=="network" else "Hierarchical Tree")
    with col_opt3:
        learning_mode = st.selectbox("Visual Learning Mode:", ["Student Mode", "Beginner Mode", "Professional Mode", "Research Mode"])
        
    if st.button("🌳 Generate / Load Visual Mind Map"):
        with st.spinner("🧠 Constructing interactive knowledge representation..."):
            results = store.collection.get(where={"video_id": selected_vid_id})
            docs = results.get("documents", [])
            full_text = " ".join(docs)
            
            chapters = []
            if 'ts_list' in st.session_state and st.session_state.ts_list:
                chapters = st.session_state.ts_list
                
            graph_data = graph_generator.extract_knowledge_graph_data(
                transcript=full_text,
                title=video_options[selected_vid_id],
                chapters=chapters,
                learning_mode=learning_mode,
                selected_model=st.session_state.selected_model
            )
            
            st.session_state.active_graph = graph_data
            st.success("Mind Map constructed successfully!")
            st.rerun()
            
    if st.session_state.active_graph:
        graph_data = st.session_state.active_graph
        html_code = graph_generator.generate_vis_html(graph_data, layout_mode=vis_layout)
        
        st.subheader("💡 Interactive Knowledge Visualization Canvas")
        st.caption("🔍 Scroll to Zoom | Drag to Pan/Re-arrange | Click Node to View Details")
        
        import streamlit.components.v1 as components
        components.html(html_code, height=600, scrolling=True)
        
        recs = graph_generator.generate_recommendations(graph_data)
        st.divider()
        st.subheader("💡 AI Learning Recommendations")
        col_rec1, col_rec2 = st.columns(2)
        with col_rec1:
            st.write("**Next Steps to Learn:**")
            for r in recs[:2]:
                st.info(f"👉 **{r['topic']}**\n*{r['reason']}*")
        with col_rec2:
            st.write("**Recommended Review Areas:**")
            for r in recs[2:]:
                st.warning(f"🎯 **{r['topic']}**\n*{r['reason']}*")
                
        st.divider()
        st.subheader("📥 Export Mind Map Structure")
        col_ex1, col_ex2, col_ex3, col_ex4 = st.columns(4)
        title_safe = video_options[selected_vid_id].replace(" ", "_")
        
        with col_ex1:
            mermaid_str = graph_generator.export_graph_mermaid(graph_data)
            st.download_button("⚙️ Mermaid Diagram (.mmd)", mermaid_str, file_name=f"{title_safe}_MindMap.mmd", mime="text/plain", key="dl_mermaid")
        with col_ex2:
            md_str = graph_generator.export_graph_markdown(graph_data)
            st.download_button("📄 Markdown Structure (.md)", md_str, file_name=f"{title_safe}_MindMap.md", mime="text/markdown", key="dl_md_graph")
        with col_ex3:
            csv_str = graph_generator.export_graph_csv(graph_data)
            st.download_button("📊 CSV Tables (Nodes/Edges)", csv_str, file_name=f"{title_safe}_MindMap.csv", mime="text/csv", key="dl_csv_graph")
        with col_ex4:
            json_str = json.dumps(graph_data, indent=2)
            st.download_button("⚙️ JSON Raw Schema", json_str, file_name=f"{title_safe}_MindMap.json", mime="application/json", key="dl_json_graph")

def create_flashcards_markdown(cards, title):
    md = f"# AI Study Flashcards: {title}\n\n"
    for idx, card in enumerate(cards):
        md += f"### Flashcard {idx+1} ({card.get('category', 'General')})\n"
        md += f"**Question:** {card['question']}\n\n"
        md += f"**Answer:** {card['answer']}\n\n"
        md += f"- **Difficulty:** {card.get('difficulty', 'Medium')} | **Chapter:** {card.get('chapter', 'General')} | **Timestamp:** {card.get('timestamp', '00:00')}\n\n"
        md += "---\n\n"
    return md

def create_flashcards_pdf(cards, title):
    pdf = BeautifulPDF(f"Flashcards: {title}")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_text_color(75, 0, 130) # Indigo
    pdf.set_font("Helvetica", 'B', 22)
    pdf.ln(60)
    pdf.cell(0, 15, "AI Study Flashcards", ln=True, align='C')
    pdf.set_font("Helvetica", 'I', 16)
    title_clean = title.encode('latin-1', 'replace').decode('latin-1')
    pdf.cell(0, 10, title_clean, ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Helvetica", size=14)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, "Generated by AI Learning Suite", ln=True, align='C')
    pdf.cell(0, 10, datetime.now().strftime('%B %d, %Y'), ln=True, align='C')
    
    pdf.add_page()
    for idx, card in enumerate(cards):
        pdf.set_font("Helvetica", 'B', 12)
        pdf.set_text_color(75, 0, 130)
        q_label = f"Card {idx+1} [{card.get('category', 'General')}] (Diff: {card.get('difficulty', 'Medium')})".encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 8, q_label, ln=True)
        pdf.set_font("Helvetica", 'B', 10)
        pdf.set_text_color(30, 30, 30)
        q_text = f"Question: {card['question']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, q_text)
        pdf.set_font("Helvetica", size=10)
        a_text = f"Answer: {card['answer']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, a_text)
        ch_text = f"Chapter: {card.get('chapter', 'General')} | Timestamp: {card.get('timestamp', '00:00')}".encode('latin-1', 'replace').decode('latin-1')
        pdf.set_font("Helvetica", 'I', 8)
        pdf.set_text_color(128, 128, 128)
        pdf.cell(0, 6, ch_text, ln=True)
        pdf.ln(4)
    return pdf.output(dest='S')

def create_flashcards_docx(cards, title):
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("AI Study Flashcards")
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_paragraph()
    
    for idx, card in enumerate(cards):
        p_title = doc.add_paragraph()
        run = p_title.add_run(f"Flashcard {idx+1} ({card.get('category', 'General')}) - {card.get('difficulty', 'Medium')}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(106, 90, 205)
        
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"Question: ")
        run_q.bold = True
        p_q.add_run(card['question'])
        
        p_a = doc.add_paragraph()
        run_a = p_a.add_run(f"Answer: ")
        run_a.bold = True
        p_a.add_run(card['answer'])
        
        p_meta = doc.add_paragraph()
        run_m = p_meta.add_run(f"Chapter: {card.get('chapter', 'General')} | Timestamp: {card.get('timestamp', '00:00')}")
        run_m.font.italic = True
        run_m.font.size = Pt(9.5)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_quiz_markdown(questions, title):
    md = f"# AI Advanced Quiz: {title}\n\n"
    for idx, q in enumerate(questions):
        md += f"### Question {idx+1} ({q.get('type', 'MCQ')} - {q.get('difficulty', 'Medium')})\n"
        md += f"**Question:** {q['question']}\n\n"
        if q.get('type') == 'MCQ' and q.get('options'):
            for opt in q['options']:
                md += f"- {opt}\n"
            md += "\n"
        md += f"**Answer:** {q['answer']}\n\n"
        md += f"**Explanation:** {q.get('explanation', '')}\n\n"
        md += f"- **Chapter:** {q.get('chapter', 'General')} | **Timestamp:** {q.get('timestamp', '00:00')}\n\n"
        md += "---\n\n"
    return md

def create_quiz_pdf(questions, title):
    pdf = BeautifulPDF(f"Quiz: {title}")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_text_color(75, 0, 130) # Indigo
    pdf.set_font("Helvetica", 'B', 22)
    pdf.ln(60)
    pdf.cell(0, 15, "AI Comprehensive Assessment Quiz", ln=True, align='C')
    pdf.set_font("Helvetica", 'I', 16)
    title_clean = title.encode('latin-1', 'replace').decode('latin-1')
    pdf.cell(0, 10, title_clean, ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Helvetica", size=14)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, "Generated by AI Learning Suite", ln=True, align='C')
    pdf.cell(0, 10, datetime.now().strftime('%B %d, %Y'), ln=True, align='C')
    
    pdf.add_page()
    for idx, q in enumerate(questions):
        pdf.set_font("Helvetica", 'B', 12)
        pdf.set_text_color(75, 0, 130)
        q_label = f"Question {idx+1} [{q.get('type', 'MCQ')}] (Diff: {q.get('difficulty', 'Medium')})".encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(0, 8, q_label, ln=True)
        
        pdf.set_font("Helvetica", 'B', 10)
        pdf.set_text_color(30, 30, 30)
        q_text = f"Question: {q['question']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, q_text)
        
        pdf.set_font("Helvetica", size=10)
        if q.get('type') == 'MCQ' and q.get('options'):
            for o in q['options']:
                o_clean = f"  - {o}".encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 6, o_clean)
                
        ans_text = f"Correct Answer: {q['answer']}".encode('latin-1', 'replace').decode('latin-1')
        pdf.set_font("Helvetica", 'B', 10)
        pdf.multi_cell(0, 6, ans_text)
        
        exp_text = f"Explanation: {q.get('explanation', '')}".encode('latin-1', 'replace').decode('latin-1')
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 6, exp_text)
        
        ch_text = f"Chapter: {q.get('chapter', 'General')} | Timestamp: {q.get('timestamp', '00:00')}".encode('latin-1', 'replace').decode('latin-1')
        pdf.set_font("Helvetica", 'I', 8)
        pdf.set_text_color(128, 128, 128)
        pdf.cell(0, 6, ch_text, ln=True)
        pdf.ln(4)
    return pdf.output(dest='S')

def create_quiz_docx(questions, title):
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("AI Assessment Quiz")
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(75, 0, 130)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_paragraph()
    
    for idx, q in enumerate(questions):
        p_title = doc.add_paragraph()
        run = p_title.add_run(f"Question {idx+1} ({q.get('type', 'MCQ')}) - {q.get('difficulty', 'Medium')}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(106, 90, 205)
        
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"Question: ")
        run_q.bold = True
        p_q.add_run(q['question'])
        
        if q.get('type') == 'MCQ' and q.get('options'):
            for o in q['options']:
                p_o = doc.add_paragraph()
                p_o.add_run(f"  o {o}")
                
        p_a = doc.add_paragraph()
        run_a = p_a.add_run(f"Answer: ")
        run_a.bold = True
        p_a.add_run(str(q['answer']))
        
        p_e = doc.add_paragraph()
        run_e = p_e.add_run(f"Explanation: ")
        run_e.bold = True
        p_e.add_run(q.get('explanation', ''))
        
        p_meta = doc.add_paragraph()
        run_m = p_meta.add_run(f"Chapter: {q.get('chapter', 'General')} | Timestamp: {q.get('timestamp', '00:00')}")
        run_m.font.italic = True
        run_m.font.size = Pt(9.5)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def flashcard_view():
    st.title("🎴 Spaced Repetition Study Flashcards")
    st.write("Retain knowledge using scientifically scheduled active-recall flashcards. Review daily to maintain your study streak!")

    storage = StudyStorage()
    store = VideoRAGStore(path="chroma_db_storage")
    indexed_videos = store.get_indexed_videos()

    if not indexed_videos:
        st.warning("⚠️ No videos indexed in the database yet. Please process a video in the Home view or get a transcript in the Notes view first.")
        return

    # User Profile Header
    profile = storage.load_user_profile()
    col_p1, col_p2, col_p3 = st.columns(3)
    with col_p1:
        st.metric("Total XP", f"🌟 {profile.get('xp', 0)} XP")
    with col_p2:
        st.metric("Daily Streak", f"🔥 {profile.get('streak', 0)} days")
    with col_p3:
        st.write("**Unlocked Badges:**")
        badges = profile.get("badges", [])
        if badges:
            st.write(", ".join([f"🏆 {b}" for b in badges]))
        else:
            st.caption("No badges unlocked yet. Keep studying!")

    st.divider()

    video_options = {v["video_id"]: v["video_title"] for v in indexed_videos}
    selected_vid_id = st.selectbox("Select Video for Flashcards:", list(video_options.keys()), format_func=lambda x: video_options[x], key="fc_select_video")

    deck = storage.load_flashcards(selected_vid_id)

    if not deck:
        st.info("💡 No flashcards generated for this video yet.")
        difficulty = st.selectbox("Select Flashcard Difficulty:", ["Beginner", "Intermediate", "Advanced"], index=1, key="fc_diff_select")
        if st.button("🎴 Generate Flashcard Deck", key="fc_generate_btn"):
            with st.spinner("🧠 Generating spaced repetition flashcards using AI..."):
                results = store.collection.get(where={"video_id": selected_vid_id})
                docs = results.get("documents", [])
                full_text = " ".join(docs)
                
                chapters = []
                if 'ts_list' in st.session_state and st.session_state.ts_list:
                    chapters = st.session_state.ts_list

                deck = study_generator.generate_flashcard_deck(
                    transcript=full_text,
                    title=video_options[selected_vid_id],
                    chapters=chapters,
                    difficulty=difficulty,
                    selected_model=st.session_state.selected_model
                )
                if deck:
                    storage.save_flashcards(selected_vid_id, deck)
                    st.success("Flashcards generated and saved successfully!")
                    st.session_state.current_card_index = 0
                    st.rerun()
                else:
                    st.error("Failed to generate flashcards deck.")
    else:
        # Progress status overview
        status_counts = {"New": 0, "Learning": 0, "Review": 0, "Mastered": 0}
        for card in deck:
            st_val = card.get("status", "New")
            if st_val in status_counts:
                status_counts[st_val] += 1
            else:
                status_counts["New"] += 1
                
        col_m1, col_m2, col_m3, col_m4 = st.columns(4)
        with col_m1:
            st.info(f"🆕 New: {status_counts['New']}")
        with col_m2:
            st.warning(f"🧠 Learning: {status_counts['Learning']}")
        with col_m3:
            st.success(f"🔄 Review: {status_counts['Review']}")
        with col_m4:
            st.success(f"🏆 Mastered: {status_counts['Mastered']}")

        # Filters
        st.divider()
        st.subheader("🎯 Filter Cards")
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            categories = ["All Categories"] + list(set(c.get("category", "Concept") for c in deck))
            selected_cat = st.selectbox("Category filter:", categories, key="fc_cat_filter")
        with col_f2:
            chapters_list = ["All Chapters"] + list(set(c.get("chapter", "General") for c in deck if c.get("chapter")))
            selected_ch = st.selectbox("Chapter filter:", chapters_list, key="fc_ch_filter")
            
        filtered_deck = []
        for card in deck:
            matches_cat = (selected_cat == "All Categories") or (card.get("category") == selected_cat)
            matches_ch = (selected_ch == "All Chapters") or (card.get("chapter") == selected_ch)
            if matches_cat and matches_ch:
                filtered_deck.append(card)
                
        if not filtered_deck:
            st.info("No flashcards match the selected filters.")
            return

        if 'current_card_index' not in st.session_state:
            st.session_state.current_card_index = 0
            
        if st.session_state.current_card_index >= len(filtered_deck):
            st.session_state.current_card_index = 0
            
        idx = st.session_state.current_card_index
        card = filtered_deck[idx]

        category = card.get('category', 'General')
        difficulty = card.get('difficulty', 'Medium')
        chapter = card.get('chapter', 'General')
        ts = card.get('timestamp', '00:00')
        importance = card.get('importance_score', 80)
        
        card_html = f"""
        <style>
        .flip-card-container {{
          perspective: 1000px;
          width: 100%;
          max-width: 600px;
          height: 320px;
          margin: 20px auto;
        }}
        .flip-card-checkbox {{
          display: none;
        }}
        .flip-card-label {{
          cursor: pointer;
          display: block;
          width: 100%;
          height: 100%;
        }}
        .flip-card-inner {{
          position: relative;
          width: 100%;
          height: 100%;
          transition: transform 0.6s;
          transform-style: preserve-3d;
          border-radius: 20px;
          box-shadow: 0 10px 30px rgba(106, 90, 205, 0.15);
        }}
        .flip-card-checkbox:checked + .flip-card-label .flip-card-inner {{
          transform: rotateY(180deg);
        }}
        .flip-card-front, .flip-card-back {{
          position: absolute;
          width: 100%;
          height: 100%;
          backface-visibility: hidden;
          -webkit-backface-visibility: hidden;
          border-radius: 20px;
          display: flex;
          flex-direction: column;
          justify-content: center;
          align-items: center;
          padding: 30px;
          box-sizing: border-box;
          border: 1px solid rgba(230, 230, 250, 0.8);
          text-align: center;
        }}
        .flip-card-front {{
          background: rgba(255, 255, 255, 0.95);
          color: #4B0082;
        }}
        .flip-card-back {{
          background: linear-gradient(135deg, #6A5ACD, #4B0082);
          color: white;
          transform: rotateY(180deg);
        }}
        .flip-card-meta {{
          font-size: 11px;
          text-transform: uppercase;
          letter-spacing: 1px;
          color: #9370DB;
          margin-bottom: 15px;
          font-weight: bold;
        }}
        .flip-card-back .flip-card-meta {{
          color: #E6E6FA;
        }}
        </style>
        <div class="flip-card-container">
          <input type="checkbox" id="card_flip_{idx}" class="flip-card-checkbox" />
          <label for="card_flip_{idx}" class="flip-card-label">
            <div class="flip-card-inner">
              <div class="flip-card-front">
                <span class="flip-card-meta">{category} • {difficulty} • Importance: {importance}%</span>
                <h3 style="margin: 0; font-family: 'Outfit', sans-serif; font-size: 20px;">{card['question']}</h3>
                <span style="font-size:12px; color:#6A5ACD; margin-top:30px; font-style: italic;">👉 Click anywhere to reveal answer</span>
              </div>
              <div class="flip-card-back">
                <span class="flip-card-meta">Answer Details ({chapter} at {ts})</span>
                <h3 style="margin: 0; font-family: 'Outfit', sans-serif; font-size: 18px; font-weight: normal; line-height: 1.5;">{card['answer']}</h3>
                <span style="font-size:12px; color:#E6E6FA; margin-top:30px; font-style: italic;">👉 Click anywhere to flip back</span>
              </div>
            </div>
          </label>
        </div>
        """
        
        st.markdown(card_html, unsafe_allow_html=True)
        
        # Navigation
        col_nav1, col_nav2, col_nav3 = st.columns([1, 2, 1])
        with col_nav1:
            if st.button("⬅️ Previous Card", key="fc_prev_btn", disabled=(idx == 0)):
                st.session_state.current_card_index -= 1
                st.rerun()
                
        with col_nav3:
            if st.button("Next Card ➡️", key="fc_next_btn", disabled=(idx == len(filtered_deck) - 1)):
                st.session_state.current_card_index += 1
                st.rerun()
                
        with col_nav2:
            st.write(f"<div style='text-align:center; font-weight:bold; margin-top:5px;'>Card {idx+1} of {len(filtered_deck)}</div>", unsafe_allow_html=True)
            
        st.divider()
        st.subheader("🎓 How did you do? (Update Spaced Repetition State)")
        col_gr1, col_gr2, col_gr3, col_gr4 = st.columns(4)
        
        # Original index mapping
        original_idx = next(i for i, c in enumerate(deck) if c['question'] == card['question'])
        
        with col_gr1:
            if st.button("🟢 Easy (Mastered)", key="gr_easy", help="Schedules review in 4 days"):
                deck[original_idx]["status"] = "Mastered"
                deck[original_idx]["interval_days"] = 4
                deck[original_idx]["reviews_count"] += 1
                storage.save_flashcards(selected_vid_id, deck)
                storage.add_xp(5)
                storage.update_streak()
                st.success("Marked Mastered! +5 XP")
                if idx < len(filtered_deck) - 1:
                    st.session_state.current_card_index += 1
                st.rerun()
                
        with col_gr2:
            if st.button("🟡 Medium (Review)", key="gr_med", help="Schedules review in 2 days"):
                deck[original_idx]["status"] = "Review"
                deck[original_idx]["interval_days"] = 2
                deck[original_idx]["reviews_count"] += 1
                storage.save_flashcards(selected_vid_id, deck)
                storage.add_xp(5)
                storage.update_streak()
                st.success("Marked Review! +5 XP")
                if idx < len(filtered_deck) - 1:
                    st.session_state.current_card_index += 1
                st.rerun()
                
        with col_gr3:
            if st.button("🔴 Hard (Learning)", key="gr_hard", help="Schedules review tomorrow"):
                deck[original_idx]["status"] = "Learning"
                deck[original_idx]["interval_days"] = 1
                deck[original_idx]["reviews_count"] += 1
                storage.save_flashcards(selected_vid_id, deck)
                storage.add_xp(5)
                storage.update_streak()
                st.success("Marked Learning! +5 XP")
                if idx < len(filtered_deck) - 1:
                    st.session_state.current_card_index += 1
                st.rerun()
                
        with col_gr4:
            is_bookmarked = card.get("bookmarked", False)
            btn_label = "⭐ Bookmarked" if is_bookmarked else "☆ Bookmark Card"
            if st.button(btn_label, key="gr_bookmark"):
                deck[original_idx]["bookmarked"] = not is_bookmarked
                storage.save_flashcards(selected_vid_id, deck)
                st.success("Bookmark toggled!")
                st.rerun()
                
        # Export panel & Reset
        st.divider()
        st.subheader("📥 Export & Reset Options")
        col_ex1, col_ex2 = st.columns(2)
        with col_ex1:
            if st.button("🔄 Reset Deck Progress", key="reset_deck_btn", help="Resets all card schedules back to 'New'"):
                for c in deck:
                    c["status"] = "New"
                    c["interval_days"] = 0
                    c["reviews_count"] = 0
                    c["bookmarked"] = False
                storage.save_flashcards(selected_vid_id, deck)
                st.success("Deck progress reset successfully!")
                st.session_state.current_card_index = 0
                st.rerun()
                
        with col_ex2:
            with st.expander("Show Deck Exporters"):
                st.write("Export this flashcard deck to study on other platforms or documents:")
                
                anki_txt = study_generator.export_anki_format(deck)
                md_txt = create_flashcards_markdown(deck, video_options[selected_vid_id])
                pdf_bytes = create_flashcards_pdf(deck, video_options[selected_vid_id])
                docx_bytes = create_flashcards_docx(deck, video_options[selected_vid_id])
                
                st.download_button("🎴 Download Anki Import Deck (.txt)", anki_txt, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_Anki.txt", mime="text/plain", key="dl_fc_anki")
                
                st.write("**Copy Paste to Quizlet:**")
                st.text_area("Copy this text and paste into Quizlet's Import utility:", anki_txt, height=120, key="quizlet_import_ta")
                
                st.download_button("📄 Download Markdown (.md)", md_txt, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_Flashcards.md", mime="text/markdown", key="dl_fc_md")
                st.download_button("📕 Download PDF Sheet", pdf_bytes, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_Flashcards.pdf", mime="application/pdf", key="dl_fc_pdf")
                st.download_button("📘 Download DOCX Word Sheet", docx_bytes, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_Flashcards.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_fc_docx")

def advanced_quiz_view():
    st.title("📝 Advanced AI Quiz & Assessment Center")
    st.write("Generate interactive custom assessments (MCQs, True/False, Fill in the Blanks), track your timing, view detailed timestamp explanations, and earn gamification badges!")

    storage = StudyStorage()
    store = VideoRAGStore(path="chroma_db_storage")
    indexed_videos = store.get_indexed_videos()

    if not indexed_videos:
        st.warning("⚠️ No videos indexed in the database yet. Please process a video in the Home view or get a transcript in the Notes view first.")
        return

    # User Profile Header
    profile = storage.load_user_profile()
    col_p1, col_p2, col_p3 = st.columns(3)
    with col_p1:
        st.metric("Total XP", f"🌟 {profile.get('xp', 0)} XP")
    with col_p2:
        st.metric("Daily Streak", f"🔥 {profile.get('streak', 0)} days")
    with col_p3:
        st.write("**Unlocked Badges:**")
        badges = profile.get("badges", [])
        if badges:
            st.write(", ".join([f"🏆 {b}" for b in badges]))
        else:
            st.caption("No badges unlocked yet. Keep studying!")

    st.divider()

    video_options = {v["video_id"]: v["video_title"] for v in indexed_videos}
    selected_vid_id = st.selectbox("Select Video for Quiz:", list(video_options.keys()), format_func=lambda x: video_options[x], key="quiz_select_video")

    if 'active_quiz_questions' not in st.session_state:
        st.session_state.active_quiz_questions = None
    if 'quiz_answers' not in st.session_state:
        st.session_state.quiz_answers = {}
    if 'quiz_submitted' not in st.session_state:
        st.session_state.quiz_submitted = False
    if 'quiz_start_time' not in st.session_state:
        st.session_state.quiz_start_time = None
    if 'quiz_duration' not in st.session_state:
        st.session_state.quiz_duration = 0
    if 'active_booster_quiz' not in st.session_state:
        st.session_state.active_booster_quiz = False
    if 'weak_topics' not in st.session_state:
        st.session_state.weak_topics = []

    if st.session_state.active_quiz_questions is None:
        st.subheader("⚙️ Quiz Settings")
        c1, c2, c3 = st.columns(3)
        with c1:
            num_q = st.selectbox("Number of Questions:", [5, 10, 15, 20], index=1, key="qz_num_q")
        with c2:
            difficulty = st.selectbox("Select Difficulty Level:", ["Easy", "Medium", "Hard", "Mixed"], index=3, key="qz_diff")
        with c3:
            quiz_type = st.selectbox("Select Question Types:", ["Mixed Quiz", "MCQ Only", "True/False Only", "Fill in the Blanks Only"], index=0, key="qz_type")

        if st.button("📝 Generate Custom Quiz", key="qz_generate_btn"):
            with st.spinner("🧠 Compiling your customized assessment with Gemini AI..."):
                results = store.collection.get(where={"video_id": selected_vid_id})
                docs = results.get("documents", [])
                full_text = " ".join(docs)
                
                chapters = []
                if 'ts_list' in st.session_state and st.session_state.ts_list:
                    chapters = st.session_state.ts_list

                questions = study_generator.generate_advanced_quiz(
                    transcript=full_text,
                    title=video_options[selected_vid_id],
                    chapters=chapters,
                    num_questions=num_q,
                    quiz_type=quiz_type,
                    difficulty=difficulty,
                    selected_model=st.session_state.selected_model
                )
                if questions:
                    st.session_state.active_quiz_questions = questions
                    st.session_state.quiz_answers = {}
                    st.session_state.quiz_submitted = False
                    st.session_state.quiz_start_time = time.time()
                    st.session_state.active_booster_quiz = False
                    st.session_state.weak_topics = []
                    st.success("Quiz generated successfully!")
                    st.rerun()
                else:
                    st.error("Failed to generate quiz.")
    else:
        questions = st.session_state.active_quiz_questions
        total_questions = len(questions)

        if not st.session_state.quiz_submitted:
            st.subheader("📝 Answer the Questions")

            total_time_allowed = total_questions * 60
            if st.session_state.quiz_start_time is None:
                st.session_state.quiz_start_time = time.time()
            elapsed_time = int(time.time() - st.session_state.quiz_start_time)
            remaining_time = total_time_allowed - elapsed_time

            if remaining_time <= 0:
                st.session_state.quiz_submitted = True
                st.session_state.quiz_duration = total_time_allowed
                st.warning("⏱️ Time is up! Your answers are being auto-submitted.")
                st.rerun()
            else:
                m = remaining_time // 60
                s = remaining_time % 60
                st.write(f"⏱️ **Time Remaining:** `{m:02d}:{s:02d}`")
                progress_pct = max(0.0, min(1.0, remaining_time / total_time_allowed))
                st.progress(progress_pct)

            for i, q in enumerate(questions):
                st.markdown(f"### Question {i+1} <span style='font-size:12px; background-color:#E6E6FA; color:#4B0082; padding:3px 8px; border-radius:10px;'>{q.get('type', 'MCQ')} • {q.get('difficulty', 'Medium')}</span>", unsafe_allow_html=True)
                st.write(f"**{q['question']}**")

                q_type = q.get("type", "MCQ")
                if q_type == "MCQ":
                    opts = q.get("options", [])
                    st.session_state.quiz_answers[i] = st.radio(f"Options for Q{i+1}:", opts, index=None, key=f"qz_ans_{i}", label_visibility="collapsed")
                elif q_type == "True/False":
                    st.session_state.quiz_answers[i] = st.radio(f"Options for Q{i+1}:", ["True", "False"], index=None, key=f"qz_ans_{i}", label_visibility="collapsed")
                elif q_type == "Fill in the Blanks":
                    st.session_state.quiz_answers[i] = st.text_input(f"Your answer for Q{i+1}:", placeholder="Type definition or term here...", key=f"qz_ans_{i}", label_visibility="collapsed")
                st.write("")

            if st.button("🏁 Submit Quiz Answers", key="qz_submit_btn"):
                st.session_state.quiz_submitted = True
                st.session_state.quiz_duration = int(time.time() - st.session_state.quiz_start_time)
                
                correct_count = 0
                weak_topics = []
                for idx, q in enumerate(questions):
                    user_ans = str(st.session_state.quiz_answers.get(idx) or "").strip().lower()
                    correct_ans = str(q.get("answer") or "").strip().lower()
                    if user_ans == correct_ans:
                        correct_count += 1
                    else:
                        ch = q.get("chapter", "General")
                        if ch not in weak_topics:
                            weak_topics.append(ch)

                st.session_state.weak_topics = weak_topics
                accuracy = int((correct_count / total_questions) * 100)
                xp_earned = correct_count * 10
                
                if accuracy == 100:
                    xp_earned += 50
                
                st.session_state.quiz_xp_awarded = xp_earned
                storage.add_xp(xp_earned)
                storage.update_streak()
                st.rerun()

        else:
            st.subheader("📊 Performance Dashboard")
            correct_count = 0
            for idx, q in enumerate(questions):
                user_ans = str(st.session_state.quiz_answers.get(idx) or "").strip().lower()
                correct_ans = str(q.get("answer") or "").strip().lower()
                if user_ans == correct_ans:
                    correct_count += 1

            accuracy = int((correct_count / total_questions) * 100)
            
            c1, c2, c3, c4 = st.columns(4)
            with c1:
                st.metric("Accuracy", f"{accuracy}%")
            with c2:
                duration_m = st.session_state.quiz_duration // 60
                duration_s = st.session_state.quiz_duration % 60
                st.metric("Time Taken", f"{duration_m}m {duration_s}s")
            with c3:
                st.metric("XP Points Gained", f"+{st.session_state.quiz_xp_awarded} XP")
            with c4:
                st.metric("Current Streak", f"🔥 {profile.get('streak', 0)} Days")

            if accuracy >= 90:
                st.success("🌟 Congratulations on your outstanding performance! You have unlocked a Mastery Certificate.")
                
                date_str = datetime.now().strftime("%B %d, %Y")
                cert_html = f"""
                <div style="
                  border: 15px solid gold;
                  padding: 30px;
                  background-color: #fbfbfb;
                  text-align: center;
                  font-family: 'Georgia', serif;
                  box-shadow: 0 10px 25px rgba(212,175,55,0.2);
                  border-radius: 10px;
                  max-width: 650px;
                  margin: 20px auto;
                  color: #333;
                ">
                  <div style="font-size: 14px; text-transform: uppercase; letter-spacing: 2px; color: #777;">Certificate of Mastery</div>
                  <div style="font-size: 40px; margin: 15px 0; color: #8B6508; font-weight: bold;">KNOWLEDGE MASTERY</div>
                  <div style="font-size: 16px; font-style: italic; margin-bottom: 25px;">This is proudly awarded to you for achieving a score of</div>
                  <div style="font-size: 48px; font-weight: bold; color: #D4AF37; margin-bottom: 25px;">{accuracy}%</div>
                  <div style="font-size: 18px; line-height: 1.5; margin-bottom: 30px;">
                    In the topic evaluation of<br>
                    <strong>{video_options[selected_vid_id]}</strong>
                  </div>
                  <div style="font-size: 12px; color: #999; border-top: 1px solid #ddd; padding-top: 20px; display: inline-block; width: 60%;">
                    Verified by YouTube AI Assistant<br>
                    Date: {date_str}
                  </div>
                </div>
                """
                st.markdown(cert_html, unsafe_allow_html=True)

            st.write("### Detailed Question Review")
            for i, q in enumerate(questions):
                user_val = st.session_state.quiz_answers.get(i) or "[No Answer]"
                correct_val = q.get("answer")
                is_correct = str(user_val).strip().lower() == str(correct_val).strip().lower()

                bg_color = "#E8F5E9" if is_correct else "#FFEBEE"
                border_color = "#2E7D32" if is_correct else "#D32F2F"
                status_icon = "✅" if is_correct else "❌"

                st.markdown(
                    f"""
                    <div style="background-color:{bg_color}; border-left:5px solid {border_color}; padding:15px; border-radius:10px; margin-bottom:15px;">
                        <strong>Question {i+1} ({q.get('type')}):</strong> {q['question']}<br>
                        <span style="font-size:13px; color:#555;">
                            {status_icon} <b>Your Answer:</b> {user_val} | <b>Correct Answer:</b> {correct_val}
                        </span>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                st.write(f"💡 **Explanation:** {q.get('explanation')}")
                st.caption(f"📌 **Citation:** Chapter: `{q.get('chapter', 'General')}` • Timestamp: `{q.get('timestamp', '00:00')}`")
                st.divider()

            st.subheader("💡 What's Next?")
            c_bt1, c_bt2 = st.columns(2)
            with c_bt1:
                if st.button("🔄 Retake Custom Quiz", key="qz_retake_btn"):
                    st.session_state.active_quiz_questions = None
                    st.session_state.quiz_answers = {}
                    st.session_state.quiz_submitted = False
                    st.rerun()

            with c_bt2:
                if st.session_state.weak_topics:
                    st.write("**Areas of improvement identified:**")
                    st.write(", ".join([f"`{topic}`" for topic in st.session_state.weak_topics]))
                    if st.button("🧠 Generate Booster Quiz for Weak Areas", key="qz_booster_btn"):
                        with st.spinner("🧠 Generating specialized booster quiz for weak areas..."):
                            results = store.collection.get(where={"video_id": selected_vid_id})
                            docs = results.get("documents", [])
                            full_text = " ".join(docs)
                            
                            booster_qs = study_generator.generate_booster_quiz(full_text, st.session_state.weak_topics, st.session_state.selected_model)
                            if booster_qs:
                                st.session_state.active_quiz_questions = booster_qs
                                st.session_state.quiz_answers = {}
                                st.session_state.quiz_submitted = False
                                st.session_state.quiz_start_time = time.time()
                                st.session_state.active_booster_quiz = True
                                st.session_state.weak_topics = []
                                st.success("Booster quiz generated!")
                                st.rerun()
                            else:
                                st.error("Failed to generate booster quiz.")

            st.divider()
            st.subheader("📥 Export & Download Quiz")
            with st.expander("Show Quiz Exporters"):
                gift_str = study_generator.export_moodle_gift(questions)
                gforms_str = study_generator.export_google_forms(questions)
                md_quiz = create_quiz_markdown(questions, video_options[selected_vid_id])
                pdf_quiz = create_quiz_pdf(questions, video_options[selected_vid_id])
                docx_quiz = create_quiz_docx(questions, video_options[selected_vid_id])

                st.download_button("📄 Download Moodle GIFT Format (.gift)", gift_str, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_Quiz.gift", mime="text/plain", key="dl_qz_gift")
                
                st.write("**Google Forms Import Schema:**")
                st.text_area("Copy this text schema for Google Forms Add-ons import:", gforms_str, height=120, key="gforms_ta")

                st.download_button("📄 Download Markdown (.md)", md_quiz, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_Quiz.md", mime="text/markdown", key="dl_qz_md")
                st.download_button("📕 Download PDF Exam Paper", pdf_quiz, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_Quiz.pdf", mime="application/pdf", key="dl_qz_pdf")
                st.download_button("📘 Download DOCX Word Exam", docx_quiz, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_Quiz.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_qz_docx")

def career_learning_view():
    st.title("💼 AI Resume Optimization & Career Intelligence")
    st.write("Convert educational video learning into resume-ready achievements, LinkedIn optimization items, and ATS keywords. Analyze skill gaps against target jobs to map your roadmap!")

    import career_storage
    from career_storage import CareerStorage
    import career_generator
    
    storage = CareerStorage()
    store = VideoRAGStore(path="chroma_db_storage")
    indexed_videos = store.get_indexed_videos()

    if not indexed_videos:
        st.warning("⚠️ No videos indexed in the database yet. Please process a video in the Home view or get a transcript in the Notes view first.")
        return

    # Portfolio stats load
    portfolio = storage.load_portfolio()
    
    # 1. Selection box of indexed videos
    video_options = {v["video_id"]: v["video_title"] for v in indexed_videos}
    selected_vid_id = st.selectbox("Select Video to Build Skill Report:", list(video_options.keys()), format_func=lambda x: video_options[x], key="cr_select_video")

    # Load Report
    report = storage.load_career_report(selected_vid_id)

    # 2. Config & Generate
    if report is None:
        st.info("💡 No resume report generated for this video yet.")
        
        c_cat, c_dur = st.columns(2)
        with c_cat:
            categories = [
                "Software Development", "Web Development", "Data Science", "Machine Learning", 
                "Deep Learning", "Artificial Intelligence", "Cloud Computing", "Cyber Security", 
                "DevOps", "Data Engineering", "Mobile Development", "Blockchain", "UI/UX Design", 
                "Business Analytics", "Project Management", "Finance", "Marketing", "Product Management", 
                "General Technical Content"
            ]
            # Try auto-detecting category if notes_category exists in session state
            default_cat_idx = 0
            if st.session_state.get("notes_category") in categories:
                default_cat_idx = categories.index(st.session_state.notes_category)
            category = st.selectbox("Select Video Domain Category:", categories, index=default_cat_idx, key="cr_domain_select")
            
        with c_dur:
            duration_mins = st.number_input("Video Duration (minutes):", min_value=1, max_value=600, value=30, key="cr_duration_input")
            
        if st.button("💼 Generate Career Intelligence Report", key="cr_generate_btn"):
            with st.spinner("🧠 Scanning transcript and mapping skill technologies with AI..."):
                results = store.collection.get(where={"video_id": selected_vid_id})
                docs = results.get("documents", [])
                full_text = " ".join(docs)
                
                chapters = []
                if 'ts_list' in st.session_state and st.session_state.ts_list:
                    chapters = st.session_state.ts_list
                    
                report = career_generator.extract_career_report(
                    transcript=full_text,
                    title=video_options[selected_vid_id],
                    chapters=chapters,
                    category=category,
                    selected_model=st.session_state.selected_model
                )
                if report:
                    storage.save_career_report(selected_vid_id, report)
                    storage.add_video_to_portfolio(
                        video_id=selected_vid_id,
                        title=video_options[selected_vid_id],
                        duration_seconds=duration_mins * 60,
                        report=report
                    )
                    st.success("Career report compiled and portfolio updated!")
                    st.rerun()
                else:
                    st.error("Failed to generate career report.")
    else:
        # Glassmorphism tabs
        t1, t2, t3, t4, t5 = st.tabs([
            "📊 Skills & Tech Stack", 
            "📄 ATS Resume & LinkedIn Optimizer", 
            "🎙️ Interview Preparation Q&A", 
            "🎯 Job Role Gap Analyzer",
            "📈 Learning Portfolio Dashboard"
        ])
        
        with t1:
            st.subheader("💡 Extracted Technical Skills")
            skills_data = report.get("skills_learned", [])
            if skills_data:
                for s in skills_data:
                    st.markdown(
                        f"""
                        <div style="background-color:#FFFFFF; padding:15px; border-radius:12px; margin-bottom:10px; border-left: 5px solid #4B0082; box-shadow: 0 4px 6px rgba(0,0,0,0.05);">
                            <h4 style="margin:0; color:#4B0082;">{s.get('name')}</h4>
                            <p style="margin:5px 0; font-size:14px;">{s.get('description')}</p>
                            <span style="font-size:11px; background-color:#E6E6FA; color:#4B0082; padding:3px 8px; border-radius:10px; font-weight:bold; margin-right:5px;">Importance: {s.get('importance_score')}%</span>
                            <span style="font-size:11px; background-color:#F3E5F5; color:#9370DB; padding:3px 8px; border-radius:10px; font-weight:bold; margin-right:5px;">Industry Demand: {s.get('demand_score')}%</span>
                            <span style="font-size:11px; background-color:#E8F5E9; color:#2E7D32; padding:3px 8px; border-radius:10px; font-weight:bold;">Resume Relevance: {s.get('resume_relevance_score')}%</span>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
            else:
                st.info("No explicit skills extracted.")
                
            st.subheader("⚙️ Technologies & Frameworks")
            tech_data = report.get("technologies", [])
            frameworks_data = report.get("frameworks", [])
            
            c_tc1, c_tc2 = st.columns(2)
            with c_tc1:
                st.write("**Technologies Stack:**")
                if tech_data:
                    for t in tech_data:
                        st.info(f"⚡ **{t.get('name')}** ({t.get('category')})\n- Context: *{t.get('usage_context')}*\n- Level: `{t.get('learning_level')}` | Mentions: {t.get('mentioned_count', 1)}")
                else:
                    st.caption("No technologies mapped.")
            with c_tc2:
                st.write("**Frameworks & Libraries:**")
                if frameworks_data:
                    for f in frameworks_data:
                        st.success(f"📦 **{f.get('name')}**\n- Purpose: *{f.get('purpose')}*\n- Demand: {f.get('industry_usage')} | Worthiness: {f.get('resume_worthiness')}%")
                else:
                    st.caption("No frameworks mapped.")
                    
            st.subheader("🛠️ Tools & General Concepts")
            tools_data = report.get("tools", [])
            concepts_data = report.get("concepts", [])
            
            c_tl1, c_tl2 = st.columns(2)
            with c_tl1:
                st.write("**Tools & Software:**")
                if tools_data:
                    for t in tools_data:
                        st.warning(f"🛠️ **{t.get('name')}** ({t.get('category')})\n- Context: *{t.get('usage_context')}*\n- Level: `{t.get('skill_level')}`")
                else:
                    st.caption("No tools mapped.")
            with c_tl2:
                st.write("**Core Concepts:**")
                if concepts_data:
                    for c in concepts_data:
                        st.write(f"- **{c.get('name')}**: {c.get('explanation')}")
                else:
                    st.caption("No concepts mapped.")
                    
        with t2:
            st.subheader("📄 Ready-to-Copy Resume Bullets")
            bullets = report.get("resume_bullets", [])
            if bullets:
                for b in bullets:
                    st.markdown(
                        f"""
                        <div style="background-color:#F9F9FB; border-left: 5px solid #6A5ACD; padding:12px; border-radius:8px; margin-bottom:10px;">
                            <span style="font-family: monospace; font-size:14px; color:#333;">{b}</span>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
            else:
                st.caption("No resume bullets generated.")
                
            st.subheader("📄 ATS Keyword Optimization")
            keywords = report.get("ats_keywords", [])
            if keywords:
                for k in keywords:
                    st.markdown(
                        f"""
                        <span style="display:inline-block; background-color:#E6E6FA; color:#4B0082; padding:5px 10px; border-radius:15px; font-weight:bold; font-size:12px; margin:5px 5px 5px 0;">
                            🏷️ {k.get('keyword')} (ATS weight: +{k.get('score_contribution')} | Demand: {k.get('demand_score')}%)
                        </span>
                        """,
                        unsafe_allow_html=True
                    )
            else:
                st.caption("No ATS keywords generated.")
                
            st.subheader("🌐 LinkedIn Profile Enhancer")
            ls = report.get("linkedin_suggestions", {})
            if ls:
                if ls.get("skills"):
                    st.write("**Suggested Skills to add:**")
                    st.write(", ".join([f"`{sk}`" for sk in ls.get("skills")]))
                if ls.get("headline_suggestions"):
                    st.write("**Suggested Headline update ideas:**")
                    for h in ls.get("headline_suggestions"):
                        st.info(f"💼 {h}")
                if ls.get("about_additions"):
                    st.write("**Suggested addition to your 'About' section summary:**")
                    st.write(ls.get("about_additions"))
                if ls.get("suggested_certifications"):
                    st.write("**Recommended Certifications to pursue:**")
                    st.write(", ".join([f"🏆 {c}" for c in ls.get("suggested_certifications")]))
            else:
                st.caption("No LinkedIn suggestions found.")

        with t3:
            st.subheader("🎙️ Conceptual & Technical Interview Prep Q&A")
            interview_qs = report.get("interview_preparation", [])
            if interview_qs:
                for idx, q in enumerate(interview_qs):
                    with st.expander(f"Q{idx+1}: {q.get('question')} (Difficulty: {q.get('difficulty', 'Medium')})"):
                        st.markdown(f"**Answer Key:**\n{q.get('answer')}")
            else:
                st.caption("No interview questions mapped.")
                
        with t4:
            st.subheader("🎯 Target Role Skill Gap Analyzer")
            st.write("Compare the skills extracted from this video (or your accumulated portfolio) against a specific target job role to find gaps and receive a learning roadmap.")
            
            target_role = st.text_input("Enter Target Job Title:", placeholder="e.g. Data Scientist, DevOps Engineer, React Developer", key="cr_target_role")
            
            if st.button("🎯 Analyze Skill Gap & Career Roadmap", key="cr_gap_btn"):
                if target_role:
                    with st.spinner(f"🧠 Querying Recruitment and Career API for {target_role}..."):
                        # Extract learned skills list
                        learned_skills_list = [s["name"] for s in report.get("skills_learned", [])]
                        
                        gap_report = career_generator.generate_skill_gap_analysis(
                            learned_skills=learned_skills_list,
                            target_role=target_role,
                            selected_model=st.session_state.selected_model
                        )
                        if gap_report:
                            st.session_state.cr_gap_report = gap_report
                            st.success("Skill gap analysis compiled successfully!")
                        else:
                            st.error("Failed to compile analysis.")
                else:
                    st.warning("Please enter a target job title first.")
                    
            if st.session_state.get("cr_gap_report"):
                gap = st.session_state.cr_gap_report
                st.divider()
                
                c_gap1, c_gap2 = st.columns([1, 2])
                with c_gap1:
                    st.metric("Career Match Score", f"{gap.get('match_score', 0)}%")
                with c_gap2:
                    st.write("**Missing Core Competencies / Skills:**")
                    missing = gap.get("missing_skills", [])
                    if missing:
                        for m in missing:
                            st.error(f"❌ **{m.get('name')}**: *{m.get('why_needed')}*")
                    else:
                        st.success("🎉 You match all core skills required for this job title!")
                        
                st.subheader("🗺️ Customized Learning Roadmap")
                for step in gap.get("learning_roadmap", []):
                    st.write(step)
                    
        with t5:
            st.subheader("📈 Accumulated Learning Portfolio")
            
            col_p1, col_p2, col_p3 = st.columns(3)
            with col_p1:
                st.metric("Videos Processed", f"🎬 {portfolio.get('videos_processed', 0)}")
            with col_p2:
                st.metric("Total Skills Mapped", f"💡 {portfolio.get('total_skills', 0)}")
            with col_p3:
                st.metric("Total Invested Hours", f"⏱️ {portfolio.get('total_hours', 0.0):.2f} hrs")
                
            st.divider()
            
            c_port1, c_port2 = st.columns(2)
            with c_port1:
                st.write("**Accumulated Technologies learned:**")
                st.write(", ".join([f"`{t}`" for t in portfolio.get("accumulated_technologies", [])]))
                
                st.write("\n**Accumulated Frameworks & Libraries:**")
                st.write(", ".join([f"`{f}`" for f in portfolio.get("accumulated_frameworks", [])]))
            with c_port2:
                st.write("**Accumulated Tools mastered:**")
                st.write(", ".join([f"`{t}`" for t in portfolio.get("accumulated_tools", [])]))
                
                st.write("\n**Accumulated Core Skills:**")
                st.write(", ".join([f"`{s}`" for s in portfolio.get("accumulated_skills", [])]))
                
            st.divider()
            st.write("**Learning & Skill Acquisition Timeline History:**")
            history = portfolio.get("skill_history", [])
            if history:
                for idx, entry in enumerate(reversed(history)):
                    with st.expander(f"🎬 [{entry.get('date')}] {entry.get('title')} ({int(entry.get('duration', 0)//60)} mins)"):
                        st.write(f"**Skills Mapped:** {', '.join(entry.get('skills', []))}")
            else:
                st.caption("No history entries recorded yet.")

        # Download Exporter Block
        st.divider()
        st.subheader("📥 Export & Download Skill Report")
        
        md_content = career_generator.create_career_markdown(report, video_options[selected_vid_id])
        pdf_bytes = career_generator.create_career_pdf(report, video_options[selected_vid_id])
        docx_bytes = career_generator.create_career_docx(report, video_options[selected_vid_id])
        json_bytes = json.dumps(report, indent=2)
        
        c_ex1, c_ex2, c_ex3, c_ex4 = st.columns(4)
        with c_ex1:
            st.download_button("📄 Download Markdown (.md)", md_content, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_CareerReport.md", mime="text/markdown", key="dl_cr_md")
        with c_ex2:
            st.download_button("📕 Download PDF Report", pdf_bytes, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_CareerReport.pdf", mime="application/pdf", key="dl_cr_pdf")
        with c_ex3:
            st.download_button("📘 Download DOCX Word", docx_bytes, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_CareerReport.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_cr_docx")
        with c_ex4:
            st.download_button("⚙️ Download JSON schema", json_bytes, file_name=f"{video_options[selected_vid_id].replace(' ', '_')}_CareerReport.json", mime="application/json", key="dl_cr_json")

def universal_export_view():
    st.title("📥 Universal Export Center")
    st.write("Manage, customize, and compile all your generated content into professionally formatted documents.")

    import document_builder
    from export_manager import ExportManager
    
    manager = ExportManager()
    store = VideoRAGStore(path="chroma_db_storage")
    indexed_videos = store.get_indexed_videos()

    if not indexed_videos:
        st.warning("⚠️ No videos indexed in the database yet. Please process a video in the Home view or get a transcript in the Notes view first.")
        return

    # Data collection simulation for demo purposes
    video_options = {v["video_id"]: v["video_title"] for v in indexed_videos}
    selected_vid_id = st.selectbox("Select Video to Export Content From:", list(video_options.keys()), format_func=lambda x: video_options[x], key="ue_select_video")
    
    video_info = {"title": video_options[selected_vid_id], "channel": "YouTube Channel"}

    t1, t2 = st.tabs(["📝 Document Builder", "📂 Download Center"])

    with t1:
        st.subheader("Compile Custom Report")
        c1, c2 = st.columns([1, 2])
        
        with c1:
            st.write("**Select Theme:**")
            theme = st.selectbox("Theme", ["Professional", "Academic", "Corporate", "Modern", "Minimal"], label_visibility="collapsed")
            st.write("**Select Format:**")
            format_choice = st.radio("Format", ["PDF", "DOCX", "TXT", "ZIP Archive"], label_visibility="collapsed")
            
        with c2:
            st.write("**Select Sections to Include:**")
            include_summary = st.checkbox("Executive Summary", value=True)
            include_notes = st.checkbox("Detailed Notes", value=True)
            include_takeaways = st.checkbox("Key Takeaways", value=True)
            include_flashcards = st.checkbox("Study Flashcards")
            include_quiz = st.checkbox("Assessment Quiz")
            include_career = st.checkbox("Career Intelligence Report")
            
        if st.button("🚀 Compile Document"):
            selected_sections = []
            data_sources = {}
            if include_summary: 
                selected_sections.append("Executive Summary")
                data_sources["Executive Summary"] = "This is a placeholder for the Executive Summary text retrieved from the database."
            if include_notes:
                selected_sections.append("Detailed Notes")
                data_sources["Detailed Notes"] = "This is a placeholder for Detailed Notes."
            if include_takeaways:
                selected_sections.append("Key Takeaways")
                data_sources["Key Takeaways"] = "- Point 1\n- Point 2\n- Point 3"
            if include_flashcards:
                selected_sections.append("Study Flashcards")
                data_sources["Study Flashcards"] = "Flashcard 1: Q / A\nFlashcard 2: Q / A"
            if include_quiz:
                selected_sections.append("Assessment Quiz")
                data_sources["Assessment Quiz"] = "Q1: A\nQ2: B"
            if include_career:
                selected_sections.append("Career Intelligence Report")
                data_sources["Career Intelligence Report"] = "Skills: Python, Data Science\nExperience: Junior"
                
            if not selected_sections:
                st.warning("Please select at least one section to include.")
            else:
                with st.spinner(f"Compiling {format_choice}..."):
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    if format_choice == "PDF":
                        file_bytes = document_builder.generate_combined_pdf(selected_sections, video_info, data_sources, theme)
                        filename = f"Export_{timestamp}.pdf"
                        file_type = "application/pdf"
                    elif format_choice == "DOCX":
                        file_bytes = document_builder.generate_combined_docx(selected_sections, video_info, data_sources, theme)
                        filename = f"Export_{timestamp}.docx"
                        file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    elif format_choice == "TXT":
                        file_bytes = document_builder.generate_combined_txt(selected_sections, video_info, data_sources)
                        filename = f"Export_{timestamp}.txt"
                        file_type = "text/plain"
                    else: # ZIP
                        # Provide TXT versions of sections in a ZIP
                        files_for_zip = [{"filename": f"{sec.replace(' ', '_')}.txt", "data": document_builder.generate_combined_txt([sec], video_info, data_sources)} for sec in selected_sections]
                        file_bytes = document_builder.generate_zip_package(files_for_zip, video_info['title'])
                        filename = f"Export_{timestamp}.zip"
                        file_type = "application/zip"
                        
                    # Save to local dir for Download Center
                    filepath = os.path.join(manager.exports_dir, filename)
                    with open(filepath, "wb") as f:
                        f.write(file_bytes)
                        
                    file_size_kb = len(file_bytes) / 1024.0
                    manager.add_export(filename, filepath, format_choice, file_size_kb, video_info["title"], selected_sections)
                    
                    st.success(f"{filename} compiled successfully! You can download it below or manage it in the Download Center.")
                    st.download_button(f"📥 Download {filename}", file_bytes, file_name=filename, mime=file_type)

    with t2:
        st.subheader("📂 Download Center & File Manifest")
        manifest = manager.load_manifest()
        
        if not manifest:
            st.info("No exports found. Compile a document in the Document Builder first.")
        else:
            for record in manifest:
                with st.expander(f"{'⭐ ' if record.get('favorite') else ''}{record['filename']} ({record.get('size_kb', 0):.1f} KB) - {record.get('video_title', 'Unknown')}"):
                    c1, c2, c3, c4 = st.columns(4)
                    
                    with c1:
                        if os.path.exists(record['filepath']):
                            with open(record['filepath'], "rb") as f:
                                st.download_button("📥 Download", f.read(), file_name=record['filename'], key=f"dl_{record['id']}")
                        else:
                            st.error("File missing from disk.")
                            
                    with c2:
                        new_name = st.text_input("Rename:", value=record['filename'], key=f"rn_{record['id']}", label_visibility="collapsed")
                        if st.button("Rename", key=f"btn_rn_{record['id']}"):
                            manager.rename_export(record['id'], new_name)
                            st.rerun()
                            
                    with c3:
                        fav_label = "Unfavorite" if record.get('favorite') else "⭐ Favorite"
                        if st.button(fav_label, key=f"btn_fav_{record['id']}"):
                            manager.toggle_favorite(record['id'])
                            st.rerun()
                            
                    with c4:
                        if st.button("🗑️ Delete", key=f"btn_del_{record['id']}"):
                            manager.delete_export(record['id'])
                            st.rerun()
                            
                    st.write("**Sections included:**", ", ".join(record.get('sections', [])))
                    st.write("**Date:**", record.get('date', 'Unknown'))
                    
            st.divider()
            st.subheader("☁️ Cloud Sync (Simulation)")
            c_g, c_o, c_d = st.columns(3)
            with c_g:
                if st.button("Sync to Google Drive"):
                    st.success("Synced to Google Drive successfully!")
            with c_o:
                if st.button("Sync to OneDrive"):
                    st.success("Synced to OneDrive successfully!")
            with c_d:
                if st.button("Sync to Dropbox"):
                    st.success("Synced to Dropbox successfully!")

def multi_video_research_view():
    st.title("📚 AI Multi-Video Summarizer & Research Center")
    st.write("Analyze multiple videos simultaneously to extract cross-video insights, detect contradictions, build comparison tables, and generate a unified knowledge base.")

    storage = MultiVideoStorage()
    
    if 'mv_transcripts' not in st.session_state:
        st.session_state.mv_transcripts = {}
    if 'mv_report' not in st.session_state:
        st.session_state.mv_report = None
        
    t_input, t_results = st.tabs(["📥 Import & Process", "📊 Multi-Video Insights"])
    
    with t_input:
        st.subheader("Input Video Sources")
        st.write("Enter up to 20 YouTube video URLs (one per line):")
        urls_input = st.text_area("Video URLs:", height=150, placeholder="https://youtube.com/watch?v=...\nhttps://youtube.com/watch?v=...")
        
        if st.button("🚀 Process & Synthesize Videos", key="btn_mv_process"):
            urls = [url.strip() for url in urls_input.split('\n') if url.strip()]
            if not urls:
                st.warning("Please enter at least one URL.")
            else:
                st.session_state.mv_transcripts = {}
                st.session_state.mv_report = None
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                transcripts_dict = {}
                for i, url in enumerate(urls):
                    status_text.text(f"Fetching transcript for video {i+1}/{len(urls)}...")
                    text = get_youtube_transcript(url)
                    if text:
                        title = f"Video {i+1}: {url[-11:]}"
                        transcripts_dict[title] = text
                    progress_bar.progress((i+1) / (len(urls)*2))
                    
                if not transcripts_dict:
                    st.error("Failed to retrieve any transcripts (make sure videos have captions).")
                else:
                    status_text.text("🤖 AI is synthesizing cross-video insights (this may take a minute)...")
                    try:
                        report = multi_video_generator.compile_multi_video_report(
                            transcripts_dict, 
                            length_mode="Detailed", 
                            selected_model=st.session_state.selected_model
                        )
                        st.session_state.mv_report = report
                        st.session_state.mv_transcripts = transcripts_dict
                        status_text.text("✅ Processing complete! View the results in the Insights tab.")
                        progress_bar.progress(1.0)
                    except Exception as e:
                        st.error(f"Error during synthesis: {e}")

    with t_results:
        report = st.session_state.mv_report
        if not report:
            st.info("No multi-video report generated yet.")
        else:
            r1, r2, r3, r4 = st.tabs(["📊 Overview & Insights", "🔀 Comparisons", "📝 Consolidated Summary", "💬 Unified RAG Chat"])
            
            with r1:
                st.subheader("💡 Common Insights (Across Videos)")
                common = report.get("insights", {}).get("common_insights", [])
                for c in common:
                    st.success(f"**{c.get('concept')}** (Mentioned in: {', '.join(c.get('videos_mentioning', []))})\n- Importance: {c.get('importance_score', 0)}/100")
                    
                st.subheader("🔍 Unique Perspectives")
                unique = report.get("insights", {}).get("unique_insights", [])
                for u in unique:
                    st.info(f"**{u.get('concept')}** (Source: {u.get('source_video')})\n- Reason: {u.get('reason_for_uniqueness')}")
                    
                st.subheader("🏆 Best Video Recommendations")
                best = report.get("best_videos", {}).get("categories", [])
                for b in best:
                    st.write(f"- **{b.get('category')}**: `{b.get('video_title')}` ({b.get('reason')})")

            with r2:
                st.subheader("⚖️ Topic Comparison Matrix")
                comp_table = report.get("comparison_table", "")
                st.markdown(comp_table)
                
                if comp_table:
                    csv_bytes = document_builder.generate_comparison_csv(comp_table)
                    st.download_button("📥 Download CSV", csv_bytes, "comparison_matrix.csv", "text/csv")
                
                st.subheader("⚠️ Detected Contradictions")
                contra = report.get("contradictions", {}).get("contradictions", [])
                if contra:
                    for ct in contra:
                        st.error(f"**Topic**: {ct.get('topic')}\n- {ct.get('video_a')}: {ct.get('statement_a')}\n- {ct.get('video_b')}: {ct.get('statement_b')}\n- **Analysis**: {ct.get('explanation')}")
                else:
                    st.success("No contradictions detected across these sources.")
                    
                st.subheader("🕸️ Content Similarity Matrix")
                sim = report.get("similarity_matrix", [])
                if sim:
                    import pandas as pd
                    titles = [s['title'] for s in sim]
                    scores = [s['scores'] for s in sim]
                    df = pd.DataFrame(scores, index=titles, columns=titles)
                    st.dataframe(df)

            with r3:
                st.subheader("📚 Master Consolidated Summary")
                st.markdown(report.get("consolidated_summary", "No summary available."))

            with r4:
                st.subheader("💬 Unified Knowledge Chat")
                st.write("Ask questions across all processed videos.")
                q = st.text_input("Ask a question:", key="mv_rag_q")
                if st.button("Ask", key="mv_rag_btn"):
                    with st.spinner("Analyzing all transcripts..."):
                        combined = " ".join([t[:5000] for t in st.session_state.mv_transcripts.values()])
                        from langchain_google_genai import ChatGoogleGenerativeAI
                        from langchain_core.messages import HumanMessage
                        model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
                        ans = model.invoke([HumanMessage(content=f"Context from multiple videos:\n{combined}\n\nQuestion: {q}\nAnswer strictly based on context. Note contradictions if they exist.")])
                        st.write(ans.content)

def upload_center_view():
    st.title("📤 Universal Media Upload Center")
    st.write("Upload local audio and video files. We will automatically transcribe them using Gemini 1.5 Pro and integrate the transcripts with all AI modules.")
    
    import local_media_processor
    from media_library import MediaLibrary
    import tempfile
    
    media_lib = MediaLibrary()
    
    uploaded_file = st.file_uploader("Upload a file (MP4, MP3, WAV, MKV, etc.)", type=["mp4", "mkv", "avi", "mov", "wmv", "webm", "mpeg", "m4v", "flv", "mp3", "wav", "flac", "aac", "m4a", "ogg", "wma"])
    
    if uploaded_file is not None:
        file_details = {"FileName": uploaded_file.name, "FileType": uploaded_file.type, "FileSize": f"{uploaded_file.size / (1024*1024):.2f} MB"}
        st.write("### File Details:")
        st.write(file_details)
        
        if st.button("🚀 Upload & Process Media"):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # 1. Save to temp
            status_text.text("Saving file locally...")
            temp_dir = "temp_media"
            os.makedirs(temp_dir, exist_ok=True)
            temp_path = os.path.join(temp_dir, uploaded_file.name)
            
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
                
            progress_bar.progress(10)
            
            # 2. Extract metadata
            status_text.text("Extracting metadata...")
            metadata = local_media_processor.extract_metadata(temp_path)
            progress_bar.progress(20)
            
            # 3. Add to library
            media_id = media_lib.add_media(uploaded_file.name, temp_path, uploaded_file.type, metadata)
            
            # 4. Extract Audio if video
            audio_path = temp_path
            if metadata.get("is_video"):
                status_text.text("Extracting audio from video (this may take a moment)...")
                extracted = local_media_processor.extract_audio_from_video(temp_path)
                if extracted:
                    audio_path = extracted
            
            progress_bar.progress(40)
            
            # 5. Transcribe using Gemini
            status_text.text("Transcribing audio using Gemini 1.5 Pro File API...")
            api_key = st.session_state.get('GEMINI_API_KEY', GEMINI_API_KEY)
            transcript = local_media_processor.transcribe_audio(audio_path, api_key=api_key)
            progress_bar.progress(80)
            
            if transcript:
                media_lib.update_transcript(media_id, transcript)
                status_text.text("Generating Summary...")
                
                # We mock process_with_gemini
                from langchain_google_genai import ChatGoogleGenerativeAI
                from langchain_core.messages import HumanMessage
                model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
                prompt = "Analyze this transcript and provide a response in Markdown:\n\nWrite an engaging executive summary.\nTranscript:\n" + transcript[:10000]
                resp = model.invoke([HumanMessage(content=prompt)])
                
                st.session_state.summary = resp.content
                st.session_state.processed_url = f"local://{media_id}"
                st.session_state.video_info = {
                    "title": uploaded_file.name,
                    "channel": "Local Upload",
                    "duration": metadata.get("duration_sec", 0),
                    "uploader": "You"
                }
                st.session_state.chat_history = []
                st.session_state.quiz_data = None
                
                # Make sure it's in RAG
                try:
                    rag_storage.index_video_in_rag(st.session_state.processed_url, uploaded_file.name, st.session_state.video_info)
                except Exception as e:
                    pass
                
                progress_bar.progress(100)
                status_text.text("✅ Processing complete!")
                st.success("Media processed successfully! You can now use all AI features from the sidebar on this file.")
                st.session_state.view = 'main'
                time.sleep(2)
                st.rerun()
            else:
                st.error("Transcription failed.")

def learning_roadmap_view():
    st.title("🗺️ AI Learning Roadmap Generator")
    st.write("Generate a personalized 4-tier skill pathway from this video.")
    
    if not st.session_state.processed_url:
        st.warning("Please process a video or upload a file first.")
        if st.button("⬅️ Go Back"):
            st.session_state.view = 'main'
            st.rerun()
        return

    import roadmap_generator
    from roadmap_storage import RoadmapStorage
    storage = RoadmapStorage()
    
    video_id = st.session_state.processed_url
    roadmap_data = storage.get_roadmap(video_id)
    
    col1, col2 = st.columns([2, 1])
    
    if roadmap_data is None:
        with col1:
            st.info("No roadmap generated for this video yet.")
            if st.button("🚀 Generate Personalized Learning Roadmap"):
                with st.spinner("AI is analyzing the content and extrapolating a complete domain roadmap..."):
                    transcript = get_youtube_transcript(st.session_state.processed_url)
                    title = st.session_state.video_info.get('title', 'Unknown Title')
                    
                    if transcript:
                        # Extract the roadmap
                        roadmap_data = roadmap_generator.generate_learning_roadmap(transcript, title, st.session_state.selected_model)
                        storage.save_roadmap(video_id, roadmap_data)
                        st.success("Roadmap generated successfully!")
                        st.rerun()
                    else:
                        st.error("Could not fetch transcript to generate roadmap.")
    else:
        # Display Roadmap
        st.subheader(f"Domain: {roadmap_data.get('domain', 'Learning Path')}")
        st.write(f"**Total Estimated Learning Time:** {roadmap_data.get('total_estimated_time', 'N/A')}")
        
        # Calculate Progress
        completed = roadmap_data.get("completed_topics", [])
        total_topics = sum([len(level.get("topics", [])) for level in roadmap_data.get("roadmap", [])])
        
        progress = len(completed) / total_topics if total_topics > 0 else 0
        st.progress(progress)
        st.write(f"**Completion Status:** {int(progress * 100)}% ({len(completed)}/{total_topics} topics learned)")
        
        for level in roadmap_data.get("roadmap", []):
            with st.expander(f"📈 {level.get('level', 'Unknown').upper()} (Time: {level.get('estimated_time', '')})"):
                st.write(f"**Focus:** {level.get('focus', '')}")
                
                st.markdown("### 📚 Core Topics")
                for topic in level.get("topics", []):
                    topic_name = topic.get("name")
                    is_done = topic_name in completed
                    
                    cols = st.columns([1, 10])
                    with cols[0]:
                        done = st.checkbox("", value=is_done, key=f"rm_chk_{topic_name}")
                        if done != is_done:
                            storage.toggle_topic_completion(video_id, topic_name, done)
                            st.rerun()
                            
                    with cols[1]:
                        st.markdown(f"**{topic_name}** - {topic.get('description')} *(Importance: {topic.get('importance')})*")
                        reqs = ", ".join(topic.get("prerequisites", []))
                        if reqs and reqs.lower() != "none":
                            st.caption(f"Prerequisites: {reqs}")
                            
                st.markdown("### 🚀 Recommended Projects")
                for proj in level.get("projects", []):
                    st.write(f"- **{proj.get('name')}**: {proj.get('description')}")
                    
                st.markdown("### 📖 Resources")
                for res in level.get("resources", []):
                    st.write(f"- {res}")
                    
        # Export logic
        st.divider()
        st.subheader("📥 Export Roadmap")
        export_col1, export_col2 = st.columns(2)
        import document_builder
        md_text = roadmap_generator.create_roadmap_markdown(roadmap_data)
        
        with export_col1:
            pdf_bytes = document_builder.create_beautiful_pdf_bytes(md_text, f"Roadmap: {roadmap_data.get('domain')}")
            if pdf_bytes:
                st.download_button("📄 Download PDF", data=pdf_bytes, file_name=f"Roadmap_{roadmap_data.get('domain')}.pdf", mime="application/pdf")
        with export_col2:
            docx_bytes = document_builder.create_docx_bytes(md_text, f"Roadmap: {roadmap_data.get('domain')}")
            if docx_bytes:
                st.download_button("📝 Download DOCX", data=docx_bytes, file_name=f"Roadmap_{roadmap_data.get('domain')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
    st.divider()
    st.subheader("🎯 Skill Gap Analysis")
    st.write("Compare your current skills against a target role to get a personalized gap roadmap.")
    
    gap_role = st.text_input("Target Job Role (e.g. Data Scientist, Cloud Engineer):")
    gap_skills = st.text_input("Your Current Skills (comma separated):")
    
    if st.button("🔍 Analyze Skill Gap"):
        if gap_role and gap_skills:
            with st.spinner("Analyzing skill gap..."):
                transcript = get_youtube_transcript(st.session_state.processed_url) or ""
                gap_data = roadmap_generator.analyze_skill_gap(gap_role, gap_skills, transcript, st.session_state.selected_model)
                
                st.write(f"### Target Role Match Score: {gap_data.get('match_score', 0)}/100")
                
                st.markdown("#### 🚨 Missing Skills")
                for ms in gap_data.get("missing_skills", []):
                    st.error(f"**{ms.get('name')}**: {ms.get('why_needed')}")
                    
                st.markdown("#### 🛠️ Bridging the Gap (Roadmap)")
                for step in gap_data.get("learning_roadmap", []):
                    st.info(step)
                    
                st.markdown("#### 🚀 Recommended Projects to Build")
                for rp in gap_data.get("recommended_projects", []):
                    st.success(f"- {rp}")
        else:
            st.warning("Please enter both target role and current skills.")

def main_view():
    input_source = st.radio("Input Source", ["YouTube URL", "Upload File"], horizontal=True)
    
    youtube_url = None
    uploaded_file = None
    
    if input_source == "YouTube URL":
        youtube_url = st.text_input("Enter YouTube Video URL", placeholder="https://www.youtube.com/watch?v=...")
    else:
        uploaded_file = st.file_uploader("Upload Audio or Video File", type=["mp3", "wav", "flac", "m4a", "aac", "ogg", "mp4", "mkv", "avi", "mov", "webm", "wmv", "mpeg"])
        
    summary_type = st.selectbox(
        "How would you like your summary?",
        ["Executive Summary (Standard)", "Bullet Points (Quick Read)", "Tweet-Length (Very Short)", "Detailed Report (Comprehensive)", "Simple Explanation (Like I'm 5)"],
        index=0
    )
    
    col1, col2 = st.columns(2)
    with col1:
        start_time = st.text_input("Start Time (optional)", placeholder="e.g. 1:20 or 00:01:20")
    with col2:
        end_time = st.text_input("End Time (optional)", placeholder="e.g. 5:45 or 00:05:45")

    if st.button("🚀 Process Media"):
        if youtube_url or uploaded_file:
            st.session_state.chat_history = []
            st.session_state.summary = None
            st.session_state.quiz_data = None
            st.session_state.quiz_submitted = False
            st.session_state.user_answers = {}
            st.session_state.view = 'main'
            
            audio_file = None
            info = {}
            source_id = None
            
            if input_source == "YouTube URL":
                st.session_state.processed_url = youtube_url
                st.session_state.is_local_file = False
                source_id = youtube_url
                audio_file, info = download_audio(youtube_url)
            else:
                from media_library import MediaLibrary
                from local_media_processor import extract_audio_from_video, extract_metadata
                import uuid
                import tempfile
                
                # File Validation
                if uploaded_file.size > 200 * 1024 * 1024:
                    st.error("File exceeds maximum size (200MB).")
                    return
                if uploaded_file.size == 0:
                    st.error("Corrupted file detected (Empty).")
                    return
                
                ext = os.path.splitext(uploaded_file.name)[1].lower()
                is_video = ext in ['.mp4', '.mkv', '.avi', '.mov', '.webm', '.wmv', '.mpeg']
                
                ml = MediaLibrary()
                file_id = str(uuid.uuid4())
                source_id = file_id
                st.session_state.processed_url = file_id
                st.session_state.is_local_file = True
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                status_text.text("Uploading File...")
                
                saved_path = ml.save_uploaded_file(file_id, uploaded_file.name, uploaded_file.getvalue(), is_video=is_video)
                st.session_state.local_file_path = saved_path
                info = extract_metadata(saved_path)
                info["title"] = uploaded_file.name
                info["uploader"] = "Local Upload"
                
                progress_bar.progress(30)
                
                if is_video:
                    status_text.text("Extracting Audio...")
                    audio_file = extract_audio_from_video(saved_path)
                else:
                    audio_file = saved_path
                    
                progress_bar.progress(60)
                status_text.text("Generating Summary...")

            if audio_file:
                st.session_state.video_info = info
                result = process_with_gemini(audio_file, start_time, end_time, summary_type)
                
                if result:
                    st.session_state.summary = result
                    
                    if input_source == "YouTube URL":
                        os.remove(audio_file)
                    elif is_video and audio_file != saved_path:
                        try:
                            os.remove(audio_file)
                        except:
                            pass
                            
                    from analytics_manager import AnalyticsManager
                    am = AnalyticsManager()
                    duration_mins = info.get('duration', 0) / 60
                    am.log_activity("Processed Media", f"Summarized: {info.get('title', 'Unknown')}", duration_minutes=duration_mins, video_id=source_id)
                    am.log_generated_content("Summary", video_id=source_id)
                    
                    index_video_in_rag(source_id, info.get('title', 'Media'), info)
                    st.rerun()

    if st.session_state.summary:
        col1, col2 = st.columns([1, 1.2])
        with col1:
            if getattr(st.session_state, 'is_local_file', False):
                ext = os.path.splitext(st.session_state.local_file_path)[1].lower()
                is_video = ext in ['.mp4', '.mkv', '.avi', '.mov', '.webm', '.wmv', '.mpeg']
                
                st.subheader("Media Preview")
                if is_video:
                    st.video(st.session_state.local_file_path)
                else:
                    st.audio(st.session_state.local_file_path)
            else:
                st.video(st.session_state.processed_url)
                
            if st.session_state.video_info:
                st.write(f"**Channel/Uploader:** {st.session_state.video_info.get('uploader', 'Unknown')}")
            
            st.divider()
            st.header("💬 Chat with this Video")
            for chat in st.session_state.chat_history:
                role = "user-message" if chat['role'] == 'user' else "ai-message"
                st.markdown(f'<div class="chat-message {role}"><b>{"You" if chat["role"]=="user" else "AI"}:</b><br>{chat["content"]}</div>', unsafe_allow_html=True)
            
            user_query = st.chat_input("Ask a question about the video...")
            if user_query:
                st.session_state.chat_history.append({"role": "user", "content": user_query})
                model = ChatGoogleGenerativeAI(model=st.session_state.selected_model)
                chat_context = genai.get_file(st.session_state.audio_file_name)
                message = HumanMessage(
                    content=[
                        {"type": "text", "text": f"Answer: {user_query}"},
                        {"type": "media", "file_uri": chat_context.uri, "mime_type": chat_context.mime_type}
                    ]
                )
                response = model.invoke([message])
                st.session_state.chat_history.append({"role": "assistant", "content": response.content})
                st.rerun()

        with col2:
            st.header("📝 AI Analysis")
            st.markdown(st.session_state.summary)
            st.divider()
            st.subheader("📥 Export & Download")
            
            if st.session_state.video_info.get('duration', 0) > 3600:
                st.success("🌟 Long Video Detected! A comprehensive, detailed PDF report has been generated for you.")
                pdf_bytes = create_pdf_bytes(st.session_state.summary, f"Detailed Report: {st.session_state.video_info.get('title', 'Video Analysis')}")
                if pdf_bytes:
                    st.download_button(label="📕 Download Comprehensive PDF Report", data=pdf_bytes, file_name=f"Comprehensive_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", mime="application/pdf")
            else:
                st.download_button(label="📄 Download Summary (.txt)", data=st.session_state.summary, file_name=f"Summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt", mime="text/plain")

def quiz_view():
    st.markdown('<div class="quiz-container">', unsafe_allow_html=True)
    st.title("🧠 Knowledge Mastery Quiz")
    st.divider()
    if st.session_state.quiz_data:
        for i, q in enumerate(st.session_state.quiz_data):
            st.subheader(f"Question {i+1}")
            st.write(f"**{q['question']}**")
            if not st.session_state.quiz_submitted:
                st.session_state.user_answers[i] = st.radio(f"Options for Q{i+1}:", q['options'], key=f"q_{i}", index=None, label_visibility="collapsed")
            else:
                selected = st.session_state.user_answers.get(i)
                correct = q['answer']
                for option in q['options']:
                    if option == correct: st.markdown(f'<div class="correct-ans">✅ {option} (Correct)</div>', unsafe_allow_html=True)
                    elif option == selected: st.markdown(f'<div class="wrong-ans">❌ {option} (Your Choice)</div>', unsafe_allow_html=True)
                    else: st.write(f"○ {option}")
            st.write("")
        st.divider()
        c1, c2, c3 = st.columns([1,1,1])
        with c1:
            if not st.session_state.quiz_submitted:
                if st.button("🔍 Reveal Answers"):
                    st.session_state.quiz_submitted = True
                    st.rerun()
            else:
                if st.button("🔄 Restart"):
                    st.session_state.quiz_submitted = False
                    st.session_state.user_answers = {}
                    st.rerun()
        with c3:
            if st.button("⬅️ Back"):
                st.session_state.view = 'main'
                st.rerun()
    else:
        st.warning("No quiz data found.")
        if st.button("⬅️ Back"):
            st.session_state.view = 'main'
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

def main():
    st.markdown(
        """
        <style>
        .stApp { background-color: #F8F9FE; background-image: radial-gradient(at 0% 0%, #E6E6FA 0, transparent 50%); }
        .block-container { background-color: rgba(255, 255, 255, 0.9); padding: 2rem 5rem !important; border-radius: 30px; backdrop-filter: blur(10px); }
        h1 { color: #4B0082 !important; font-weight: 800; }
        .stButton>button { background: linear-gradient(135deg, #9370DB, #6A5ACD); color: white; border-radius: 15px; width: 100%; }
        .chat-message { padding: 1rem; border-radius: 15px; margin-bottom: 1rem; border: 1px solid #E6E6FA; }
        .user-message { background-color: #F3E5F5; border-left: 5px solid #9370DB; }
        .ai-message { background-color: #FFFFFF; border-left: 5px solid #6A5ACD; }
        .quiz-container { background-color: #FFFFFF; padding: 40px; border-radius: 30px; box-shadow: 0 20px 50px rgba(147, 112, 219, 0.1); }
        .correct-ans { color: #2E7D32; font-weight: bold; background-color: #E8F5E9; padding: 10px; border-radius: 10px; margin: 5px 0; }
        .wrong-ans { color: #D32F2F; font-weight: bold; background-color: #FFEBEE; padding: 10px; border-radius: 10px; margin: 5px 0; }
        </style>
        """,
        unsafe_allow_html=True
    )

    with st.sidebar:
        st.header("Navigation")
        if st.button("🏠 Home"):
            st.session_state.view = 'main'
            st.rerun()
            
        if st.button("📊 My Dashboard"):
            st.session_state.view = 'dashboard'
            st.rerun()

            
        if st.button("📝 Notes Generator"):
            st.session_state.view = 'notes_gen'
            st.rerun()
            
        if st.button("💡 Key Takeaways"):
            st.session_state.view = 'takeaways'
            st.rerun()
            
        if st.button("⏰ Timestamp Summary"):
            st.session_state.view = 'timestamp'
            st.rerun()
            
        if st.button("💬 RAG Chatbot"):
            st.session_state.view = 'rag_chat'
            st.rerun()
            
        if st.button("🌳 Mind Map Generator"):
            st.session_state.view = 'mind_map'
            st.rerun()
            
        if st.button("🎴 Flashcards"):
            st.session_state.view = 'flashcards'
            st.rerun()
            
        if st.button("📝 Advanced Quiz"):
            st.session_state.view = 'advanced_quiz'
            st.rerun()
            
        if st.button("💼 Resume Learning"):
            st.session_state.view = 'resume_learning'
            st.rerun()
            
        if st.button("📥 Export Center"):
            st.session_state.view = 'export_center'
            st.rerun()
            
        if st.button("📤 Media Upload Center"):
            st.session_state.view = 'media_upload'
            st.rerun()
            
        if st.button("🗺️ Learning Roadmap"):
            st.session_state.view = 'learning_roadmap'
            st.rerun()
            
        if st.button("📚 Multi-Video Center"):
            st.session_state.view = 'multi_video_center'
            st.rerun()
            
        st.divider()
        st.subheader("⚙️ AI Engine Settings")
        st.session_state.selected_model = st.selectbox(
            "Select AI Model (Switch if Quota hit):",
            ["gemini-pro-latest", "gemini-flash-latest", "gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.0-flash", "gemini-flash-lite-latest"],
            index=0
        )
        st.info("💡 Tip: If you hit a limit (429 error), just switch to a different model name above!")

        if st.session_state.summary:
            st.divider()
            if st.button("📝 Learning Quiz"):
                if not st.session_state.quiz_data:
                    st.session_state.quiz_data = generate_quiz()
                st.session_state.view = 'quiz'
                st.rerun()
        st.divider()
        if st.button("🗑️ Clear All"):
            for key in list(st.session_state.keys()): del st.session_state[key]
            st.rerun()

    if st.session_state.view == 'quiz':
        quiz_view()
    elif st.session_state.view == 'notes_gen':
        notes_generator_view()
    elif st.session_state.view == 'takeaways':
        takeaways_generator_view()
    elif st.session_state.view == 'timestamp':
        timestamp_summary_view()
    elif st.session_state.view == 'rag_chat':
        rag_chatbot_view()
    elif st.session_state.view == 'mind_map':
        mind_map_view()
    elif st.session_state.view == 'flashcards':
        flashcard_view()
    elif st.session_state.view == 'advanced_quiz':
        advanced_quiz_view()
    elif st.session_state.view == 'resume_learning':
        career_learning_view()
    elif st.session_state.view == 'export_center':
        universal_export_view()
    elif st.session_state.view == 'multi_video_center':
        multi_video_research_view()
    elif st.session_state.view == 'media_upload':
        upload_center_view()
    elif st.session_state.view == 'learning_roadmap':
        learning_roadmap_view()
    elif st.session_state.view == 'dashboard':
        import dashboard_view
        dashboard_view.render_dashboard()
    else:
        st.title("YouTube AI Assistant ✨")
        main_view()

if __name__ == "__main__":
    main()
